#!/usr/bin/env python3
"""Generate an exit ticket (.docx or .txt) from a lesson plan and optional Haiku-generated content.

Exit Ticket Content Schema (JSON):
  {
    "learning_intention": "override string (optional)",
    "questions": [
      "What is the primary function of the mitochondria?",
      "Explain one example of how X leads to Y.",
      "Extension: design an experiment to test Z."
    ],
    "metacognitive": "Rate your confidence 1–5 and name one thing still confusing."
  }

Usage:
  python generate_exit_ticket.py \
    --plan path/to/plan.md \
    --date 2026-04-22 \
    --subject chem \
    --output ~/outputs/2026-04-22_chem_exit-ticket.docx \
    [--template path/to/exit-ticket-template.docx] \
    [--content-file path/to/content.json] \
    [--content-json '{"questions": [...]}'] \
    [--config ~/config.yaml] \
    [--allow-names "Maria,John"]

Use a `.txt` output path for Google Forms-ready plain text.
Requires python-docx for `.docx` output (see requirements.txt).
"""
from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path
from typing import Optional, Any

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Import shared PII scanner + allowlist resolver. When run as a script,
# the containing dir is already on sys.path; when imported as a module,
# we still prefer the sibling pii_scan.py. A broken/incomplete install
# must abort loudly — no fallback stubs, since those would persist
# student names to disk.
_HERE = Path(__file__).resolve().parent
if str(_HERE) not in sys.path:
    sys.path.insert(0, str(_HERE))
from artifact_common import load_config, resolve_output_path
from pii_scan import (
    scan_for_pii,
    scan_docx_for_pii,
    resolve_allowlist,
    load_plan_md,
    replace_docx_placeholders,
    find_unresolved_docx_placeholders,
)
from content_schema import (
    ContentValidationError,
    load_and_validate_content,
    validate_exit_ticket_content,
)
from plan_parser import parse_plan


def extract_day_fields(md: str, target_date: str) -> Optional[dict]:
    """Pull the target date's sections from a weekly markdown plan."""
    week = parse_plan(md)
    for day in week.days:
        if day.date != target_date:
            continue
        return {
            "date": day.date,
            "day_name": day.day_name,
            "learning_intention": day.learning_intention,
            "success_criteria": list(day.success_criteria),
            "subject": week.subject,
        }
    return None


def _build_from_scratch(fields: dict, content: dict):
    """Build an exit ticket Document with Haiku-provided content or blanks.

    Returns the in-memory Document so the caller can scan before saving.
    """
    doc = Document()

    title = doc.add_heading("Exit Ticket", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle_text = f"{fields['date']}"
    if fields.get("subject"):
        subtitle_text += f" — {fields['subject']}"
    subtitle = doc.add_paragraph(subtitle_text)
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.runs[0].font.size = Pt(11)

    doc.add_paragraph()

    learning_intention = content.get("learning_intention") or fields.get("learning_intention", "")
    if learning_intention:
        doc.add_heading("Learning Intention", level=2)
        doc.add_paragraph(learning_intention, style="Normal")
        doc.add_paragraph()

    doc.add_heading("Check for Understanding", level=2)

    questions = content.get("questions", []) or []
    if questions:
        for i, question in enumerate(questions, 1):
            doc.add_heading(f"Question {i}", level=3)
            doc.add_paragraph(question)
            doc.add_paragraph("Student response:", style="Normal")
            doc.add_paragraph()
    else:
        # No questions supplied — leave 2 blank slots so the doc is still usable.
        for i in (1, 2):
            doc.add_heading(f"Question {i}", level=3)
            doc.add_paragraph("[blank space for response]")
            doc.add_paragraph()

    doc.add_heading("Reflection", level=2)
    metacognitive = content.get("metacognitive", "Rate your confidence 1–5 and name one thing still confusing.")
    doc.add_paragraph(metacognitive)
    doc.add_paragraph("Student response:", style="Normal")

    return doc


def _count_question_placeholders(doc) -> int:
    """Return the highest N for which {{QUESTION_N}} appears in the template."""
    import re as _re
    highest = 0
    pat = _re.compile(r"\{\{QUESTION_(\d+)\}\}")

    def _scan(text):
        nonlocal highest
        for m in pat.finditer(text):
            n = int(m.group(1))
            if n > highest:
                highest = n

    for para in doc.paragraphs:
        _scan(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _scan(para.text)
    return highest


def _build_from_template(fields: dict, content: dict, template_path: Path):
    """Fill template placeholders with content. Returns the Document.

    Supports ``{{QUESTION_N}}`` for N up to 10. Any questions that don't
    have a matching placeholder in the template are appended as an
    "Additional Questions" section so they aren't silently dropped.
    """
    doc = Document(str(template_path))

    questions = content.get("questions", []) or []
    placeholder_ceiling = _count_question_placeholders(doc)

    replacements = {
        "{{DATE}}": fields["date"],
        "{{SUBJECT}}": fields.get("subject", ""),
        "{{LEARNING_INTENTION}}": content.get("learning_intention") or fields.get("learning_intention", ""),
    }
    for i in range(1, 11):
        replacements[f"{{{{QUESTION_{i}}}}}"] = questions[i - 1] if i - 1 < len(questions) else ""

    replace_docx_placeholders(doc, replacements)

    # Overflow: any questions the template couldn't accommodate get appended
    # at the end so nothing is dropped.
    if placeholder_ceiling and len(questions) > placeholder_ceiling:
        doc.add_paragraph()
        doc.add_heading("Additional Questions", level=2)
        for i, question in enumerate(questions[placeholder_ceiling:],
                                     start=placeholder_ceiling + 1):
            doc.add_heading(f"Question {i}", level=3)
            doc.add_paragraph(question)
            doc.add_paragraph("Student response:", style="Normal")
            doc.add_paragraph()

    return doc


def _build_forms_text(fields: dict, content: dict) -> str:
    """Build copy-paste text for a Google Forms-style exit ticket."""
    lines = ["Exit Ticket", f"Date: {fields['date']}"]
    if fields.get("subject"):
        lines.append(f"Subject: {fields['subject']}")

    learning_intention = content.get("learning_intention") or fields.get(
        "learning_intention", ""
    )
    if learning_intention:
        lines.extend(["", "Learning Intention", learning_intention])

    questions = content.get("questions", []) or []
    lines.extend(["", "Questions"])
    if questions:
        for i, question in enumerate(questions, 1):
            lines.append(f"{i}. {question}")
    else:
        lines.extend([
            "1. [Add a question aligned to today's success criteria.]",
            "2. [Add a second check-for-understanding question.]",
        ])

    lines.extend([
        "",
        "Reflection",
        content.get(
            "metacognitive",
            "Rate your confidence 1-5 and name one thing still confusing.",
        ),
    ])
    return "\n".join(_escape_formula_line(line) for line in lines) + "\n"


def _escape_formula_line(line: str) -> str:
    stripped = line.lstrip()
    if stripped[:1] not in "=+-@":
        return line
    leading_ws = line[: len(line) - len(stripped)]
    return f"{leading_ws}'{stripped}"


def main() -> int:
    ap = argparse.ArgumentParser(description="Generate exit ticket from lesson plan.")
    ap.add_argument("--plan", required=True, help="Path to markdown or .docx lesson plan")
    ap.add_argument("--date", required=True, help="Target date (YYYY-MM-DD)")
    ap.add_argument("--subject", help="Subject id (for display)")
    ap.add_argument("--output", required=True, help="Output .docx path")
    ap.add_argument("--template", help="Optional .docx template to use as base")
    ap.add_argument("--content-file", help="Path to Haiku-generated content JSON")
    ap.add_argument("--content-json", help="Inline Haiku-generated content JSON")
    ap.add_argument("--config", help="Config file for PII allowlist (defaults to ~/Documents/Lesson Plan Magic/config.yaml)")
    ap.add_argument(
        "--allow-names",
        help=(
            "Comma-separated bare names to allow "
            "(teacher/public figures only; keyword-shaped PII still fails)"
        ),
    )
    ap.add_argument("--allow-anywhere", action="store_true", help="Allow writing outside ~/Documents/Lesson Plan Magic/outputs/")
    args = ap.parse_args()

    plan_path = Path(args.plan)
    try:
        md = load_plan_md(plan_path)
    except (FileNotFoundError, ValueError) as e:
        print(f"Error: {e}", file=sys.stderr)
        return 1

    try:
        fields = extract_day_fields(md, args.date)
    except ValueError as e:
        print(f"Error: {e}", file=sys.stderr)
        return 1
    if fields is None:
        print(f"Error: date {args.date} not found in plan.", file=sys.stderr)
        return 1

    if args.subject:
        fields["subject"] = args.subject

    try:
        content = load_and_validate_content(
            content_file=args.content_file,
            content_json=args.content_json,
            validator=validate_exit_ticket_content,
        )
    except ContentValidationError as e:
        print(f"Error: invalid content JSON: {e}", file=sys.stderr)
        return 1

    try:
        cfg = load_config(args.config)
    except (FileNotFoundError, RuntimeError, ValueError) as e:
        print(f"Error: {e}", file=sys.stderr)
        return 1
    teacher_name, approved_names = resolve_allowlist(
        md, cfg, cli_allow_names=args.allow_names or "",
    )

    merged_text = md + "\n" + json.dumps(content, default=str)
    pii_error = scan_for_pii(merged_text, teacher_name=teacher_name, approved_names=approved_names)
    if pii_error:
        print(
            "Error: PII check failed. Remove student-identifying or sensitive information and retry.",
            file=sys.stderr,
        )
        return 1

    try:
        output_path = resolve_output_path(
            args.output, allow_anywhere=args.allow_anywhere
        )
    except ValueError as e:
        print(f"Error: {e}", file=sys.stderr)
        return 1
    suffix = output_path.suffix.lower()
    text_mode = suffix == ".txt"
    if suffix not in ("", ".docx", ".txt"):
        print(
            f"Error: unsupported output extension '{suffix}'. Use .docx or .txt.",
            file=sys.stderr,
        )
        return 1

    if text_mode:
        if args.template:
            print(
                "Error: templates are only supported for .docx exit tickets.",
                file=sys.stderr,
            )
            return 1
        rendered = _build_forms_text(fields, content)
        pii_error = scan_for_pii(
            rendered,
            teacher_name=teacher_name,
            approved_names=approved_names,
        )
        if pii_error:
            print(
                "Error: PII check failed on rendered ticket. Remove student-identifying or sensitive information and retry.",
                file=sys.stderr,
            )
            return 1
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text(rendered, encoding="utf-8")
        print(f"Wrote: {output_path}")
        return 0

    mode_note = ""
    if args.template:
        tpl = Path(args.template)
        if tpl.exists():
            doc = _build_from_template(fields, content, tpl)
            mode_note = " (from template)"
        else:
            print(f"Warning: template not found ({tpl}), generating from scratch.", file=sys.stderr)
            doc = _build_from_scratch(fields, content)
    else:
        doc = _build_from_scratch(fields, content)

    unresolved = find_unresolved_docx_placeholders(doc)
    if unresolved:
        print(
            "Error: unresolved template placeholders remain after fill: "
            + ", ".join(unresolved[:10]),
            file=sys.stderr,
        )
        return 1

    pii_error = scan_docx_for_pii(
        doc, teacher_name=teacher_name, approved_names=approved_names,
    )
    if pii_error:
        print(
            "Error: PII check failed on rendered ticket. Remove student-identifying or sensitive information and retry.",
            file=sys.stderr,
        )
        return 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Wrote{mode_note}: {output_path}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
