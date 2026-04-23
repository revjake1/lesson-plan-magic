#!/usr/bin/env python3
"""Generate a do-now / bell ringer (.docx) from a lesson plan and optional Haiku-generated content.

Do-Now Content Schema (JSON):
  {
    "prompt": "Sketch what happened yesterday when we mixed the two solutions...",
    "context": "Reviews the lab you ran Monday.",
    "instructions": "5 minutes. Work alone. Share your one-sentence answer.",
    "period": "3rd"
  }

Usage:
  python generate_do_now.py \
    --plan path/to/plan.md \
    --date 2026-04-22 \
    --subject chem \
    --output ~/outputs/2026-04-22_chem_do-now.docx \
    [--template path/to/do-now-strip.docx] \
    [--content-file path/to/content.json] \
    [--content-json '{"prompt": "..."}'] \
    [--config ~/config.yaml] \
    [--allow-names "Maria,John"]
    [--allow-anywhere]

Requires python-docx (see requirements.txt).
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path
from typing import Optional, Any

_SHARED_DIR = Path(__file__).resolve().parents[3] / "shared"
if str(_SHARED_DIR) not in sys.path:
    sys.path.insert(0, str(_SHARED_DIR))

try:
    from runtime_bootstrap import ensure_plugin_runtime_or_exit
except ImportError:  # pragma: no cover - exercised by isolated script tests
    ensure_plugin_runtime_or_exit = None

if ensure_plugin_runtime_or_exit is not None:
    ensure_plugin_runtime_or_exit(__file__)

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

_HERE = Path(__file__).resolve().parent
if str(_HERE) not in sys.path:
    sys.path.insert(0, str(_HERE))
from artifact_common import load_config, load_structured_day, resolve_output_path
# Hard dependency: pii_scan carries the PII checks and the plan
# loader. A broken/incomplete install must abort loudly, not fall back to
# no-op stubs that would persist student names to disk.
from pii_scan import (
    scan_for_pii,
    scan_docx_for_pii,
    resolve_allowlist,
    load_plan_md,
    replace_docx_placeholders,
    find_unresolved_docx_placeholders,
)
from content_schema import (
    ContentValidationError,
    load_and_validate_content,
    validate_do_now_content,
)


def extract_day_fields(
    plan_path: Path, target_date: str, *, plan_md: str
) -> Optional[dict]:
    """Pull the target date's sections from structured plan data."""
    day = load_structured_day(plan_path, target_date, plan_md=plan_md)
    if day is None:
        return None
    return {
        "date": day["date"],
        "day_name": day["day_name"],
        "learning_intention": day["learning_intention"],
        "subject": day["subject"],
    }


def _build_from_scratch(fields: dict, content: dict):
    """Build a do-now Document with Haiku-provided content or blanks.

    Returns the in-memory Document so the caller can scan before saving.
    """
    doc = Document()

    title = doc.add_heading("Do Now", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle_text = f"{fields['date']}"
    if fields.get("subject"):
        subtitle_text += f" — {fields['subject']}"
    if content.get("period"):
        subtitle_text += f" — {content['period']}"
    subtitle = doc.add_paragraph(subtitle_text)
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.runs[0].font.size = Pt(11)

    doc.add_paragraph()

    doc.add_heading("Opening Prompt", level=2)

    prompt = content.get("prompt")
    if prompt:
        doc.add_paragraph(prompt)
    else:
        doc.add_paragraph("[blank space for prompt]")
    if content.get("context"):
        doc.add_paragraph(content["context"])

    doc.add_paragraph()
    doc.add_paragraph("Your response:", style="Normal")
    doc.add_paragraph()

    doc.add_heading("Instructions", level=3)
    instructions = content.get("instructions", "You have 5 minutes. Write or draw your thinking. Be ready to share one idea with the class.")
    inst_para = doc.add_paragraph(instructions)
    inst_para.runs[0].font.size = Pt(11)

    return doc


def _build_from_template(fields: dict, content: dict, template_path: Path):
    """Fill template placeholders with content. Returns the Document."""
    doc = Document(str(template_path))

    replacements = {
        "{{DATE}}": fields["date"],
        "{{SUBJECT}}": fields.get("subject", ""),
        "{{PERIOD}}": content.get("period", ""),
        "{{DO_NOW_PROMPT}}": content.get("prompt", ""),
    }

    replace_docx_placeholders(doc, replacements)

    return doc


def main() -> int:
    ap = argparse.ArgumentParser(description="Generate do-now from lesson plan.")
    ap.add_argument("--plan", required=True, help="Path to markdown or .docx lesson plan")
    ap.add_argument("--date", required=True, help="Target date (YYYY-MM-DD)")
    ap.add_argument("--subject", help="Subject id (for display)")
    ap.add_argument("--output", required=True, help="Output .docx path")
    ap.add_argument("--template", help="Optional .docx template to use as base")
    ap.add_argument("--content-file", help="Path to Haiku-generated content JSON")
    ap.add_argument("--content-json", help="Inline Haiku-generated content JSON")
    ap.add_argument("--config", help="Config file for PII allowlist")
    ap.add_argument(
        "--allow-names",
        help=(
            "Comma-separated bare names to allow "
            "(teacher/public figures only; keyword-shaped PII still fails)"
        ),
    )
    ap.add_argument("--allow-anywhere", action="store_true", help="Allow writing outside ~/Documents/Lesson Plan Magic/outputs/")
    args = ap.parse_args()

    plan_path = Path(args.plan)
    try:
        md = load_plan_md(plan_path)
    except (FileNotFoundError, ValueError) as e:
        print(f"Error: {e}", file=sys.stderr)
        return 1

    try:
        fields = extract_day_fields(plan_path, args.date, plan_md=md)
    except ValueError as e:
        print(f"Error: {e}", file=sys.stderr)
        return 1
    if fields is None:
        print(f"Error: date {args.date} not found in plan.", file=sys.stderr)
        return 1

    if args.subject:
        fields["subject"] = args.subject

    try:
        content = load_and_validate_content(
            content_file=args.content_file,
            content_json=args.content_json,
            validator=validate_do_now_content,
        )
    except ContentValidationError as e:
        print(f"Error: invalid content JSON: {e}", file=sys.stderr)
        return 1

    try:
        cfg = load_config(args.config)
    except (FileNotFoundError, RuntimeError, ValueError) as e:
        print(f"Error: {e}", file=sys.stderr)
        return 1
    teacher_name, approved_names = resolve_allowlist(
        md, cfg, cli_allow_names=args.allow_names or "",
    )

    merged_text = md + "\n" + json.dumps(content, default=str)
    pii_error = scan_for_pii(merged_text, teacher_name=teacher_name, approved_names=approved_names)
    if pii_error:
        print(
            "Error: PII check failed. Remove student-identifying or sensitive information and retry.",
            file=sys.stderr,
        )
        return 1

    try:
        output_path = resolve_output_path(
            args.output, allow_anywhere=args.allow_anywhere
        )
    except ValueError as e:
        print(f"Error: {e}", file=sys.stderr)
        return 1

    mode_note = ""
    if args.template:
        tpl = Path(args.template)
        if tpl.exists():
            doc = _build_from_template(fields, content, tpl)
            mode_note = " (from template)"
        else:
            print(f"Warning: template not found ({tpl}), generating from scratch.", file=sys.stderr)
            doc = _build_from_scratch(fields, content)
    else:
        doc = _build_from_scratch(fields, content)

    unresolved = find_unresolved_docx_placeholders(doc)
    if unresolved:
        print(
            "Error: unresolved template placeholders remain after fill: "
            + ", ".join(unresolved[:10]),
            file=sys.stderr,
        )
        return 1

    pii_error = scan_docx_for_pii(
        doc, teacher_name=teacher_name, approved_names=approved_names,
    )
    if pii_error:
        print(
            "Error: PII check failed on rendered do-now. Remove student-identifying or sensitive information and retry.",
            file=sys.stderr,
        )
        return 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Wrote{mode_note}: {output_path}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
