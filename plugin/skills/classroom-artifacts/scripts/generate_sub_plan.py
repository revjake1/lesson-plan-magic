#!/usr/bin/env python3
"""Generate a sub-ready plan (.docx) from a lesson plan and optional Haiku-generated content.

Sub Plan Content Schema (JSON):
  {
    "learning_focus": "Students practice independent application of Newton's second law.",
    "periods": "1st, 3rd, 6th",
    "activity": {
      "steps": ["Hand out worksheet...", "Read directions aloud...", "Students work..."],
      "estimated_min": 40
    },
    "materials": ["Worksheet: 'Forces Practice Set' — top drawer", "Calculators — cart by door"],
    "backup": "If students finish early: read the next textbook section...",
    "emergency": {
      "nearby_teacher": "Jordan Lee (fictional), Room 204",
      "office_extension": "x201"
    },
    "return_notes": "Please leave any completed worksheets on my desk..."
  }

All names and room numbers in this schema example are fictional placeholders.

Usage:
  python generate_sub_plan.py \
    --plan path/to/plan.md \
    --date 2026-04-22 \
    --subject chem \
    --output ~/outputs/2026-04-22_chem_sub-plan.docx \
    [--template path/to/sub-plan-template.docx] \
    [--content-file path/to/content.json] \
    [--content-json '{"learning_focus": "..."}'] \
    [--period "3rd Period"] \
    [--teacher-name "Mr. Smith"] \
    [--config ~/config.yaml] \
    [--allow-names "Maria,John"]
    [--allow-anywhere]

Requires python-docx (see requirements.txt).
"""
from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path
from typing import Optional, Any

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

_HERE = Path(__file__).resolve().parent
if str(_HERE) not in sys.path:
    sys.path.insert(0, str(_HERE))
from artifact_common import load_config, resolve_output_path
# Hard dependency: pii_scan carries the PII checks and the plan
# loader. A broken/incomplete install must abort loudly, not fall back to
# no-op stubs that would persist student names to disk.
from pii_scan import (
    scan_for_pii,
    scan_docx_for_pii,
    resolve_allowlist,
    load_plan_md,
    replace_docx_placeholders,
    find_unresolved_docx_placeholders,
)
from content_schema import (
    ContentValidationError,
    load_and_validate_content,
    validate_sub_plan_content,
)
from plan_parser import parse_plan


def extract_day_fields(md: str, target_date: str) -> Optional[dict]:
    """Pull the target date's sections from a weekly markdown plan."""
    week = parse_plan(md)
    for day in week.days:
        if day.date != target_date:
            continue
        return {
            "date": day.date,
            "day_name": day.day_name,
            "learning_intention": day.learning_intention,
            "agenda": list(day.agenda),
            "materials": list(day.materials),
            "subject": week.subject,
        }
    return None


def _activity_lines(fields: dict, content: dict) -> list[str]:
    """Build the activity block, falling back to the source plan when needed."""
    lines: list[str] = []

    learning_focus = (
        content.get("learning_focus") or fields.get("learning_intention", "")
    ).strip()
    if learning_focus:
        lines.append(f"Learning Focus: {learning_focus}")

    activity_steps = (
        content.get("activity", {}).get("steps") or fields.get("agenda", [])
    )
    for step in activity_steps:
        lines.append(f"• {step}")

    return lines


def _format_emergency_text(content: dict) -> str:
    """Render emergency contact details into the starter-template line shape."""
    emergency = content.get("emergency", {}) or {}
    office_extension = (emergency.get("office_extension") or "").strip()
    nearby_teacher = (emergency.get("nearby_teacher") or "").strip()
    if not office_extension and not nearby_teacher:
        return ""
    return (
        f"Front office ext. {office_extension or '____'}   ·   "
        f"Nearest colleague: {nearby_teacher or 'Room ____'}   ·   "
        "Department chair: ____"
    )


def _build_from_scratch(fields: dict, content: dict,
                       period: Optional[str] = None,
                       teacher_name: Optional[str] = None):
    """Build a sub-ready plan Document. Returns it so the caller can scan/save."""
    doc = Document()

    header = doc.add_heading("Substitute Teacher Plan", level=1)
    header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    info_para = doc.add_paragraph()
    info_para.add_run(f"Date: {fields['date']}  |  ").bold = False
    info_para.add_run(f"Subject: {fields.get('subject', 'N/A')}").bold = False
    if period or content.get("periods"):
        periods_str = content.get("periods") or period or ""
        info_para.add_run(f"  |  Period: {periods_str}").bold = False
    info_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    if teacher_name:
        doc.add_paragraph(f"Teacher: {teacher_name}")

    doc.add_paragraph()

    doc.add_heading("Attendance & Classroom Setup", level=2)
    att_note = doc.add_paragraph("[blank space for attendance notes]")
    doc.add_paragraph()

    doc.add_heading("Today's Main Activity", level=2)

    learning_focus = content.get("learning_focus") or fields.get("learning_intention", "")
    if learning_focus:
        doc.add_paragraph(f"Learning Focus: {learning_focus}")

    doc.add_heading("Activity Instructions:", level=3)
    activity_steps = content.get("activity", {}).get("steps") or fields.get("agenda", [])
    if activity_steps:
        for step in activity_steps:
            doc.add_paragraph(step, style="List Bullet")
        estimated_min = content.get("activity", {}).get("estimated_min")
        if estimated_min:
            doc.add_paragraph(f"Estimated time: {estimated_min} minutes")
    else:
        doc.add_paragraph("[blank space for activity steps]")

    doc.add_paragraph()

    doc.add_heading("Materials & Location", level=2)
    materials_list = content.get("materials") or fields.get("materials", [])
    if materials_list:
        for material in materials_list:
            doc.add_paragraph(material, style="List Bullet")
    else:
        doc.add_paragraph("[blank space for materials list]")

    doc.add_paragraph()

    doc.add_heading("If Students Finish Early (Backup Plan)", level=2)
    backup = content.get("backup", "[blank space for backup activity]")
    doc.add_paragraph(backup)

    doc.add_paragraph()

    doc.add_heading("If There's a Problem", level=2)
    emergency = content.get("emergency", {})
    escalation = doc.add_paragraph()
    escalation.add_run("Nearby teacher: ").bold = True
    escalation.add_run(f"{emergency.get('nearby_teacher', '[blank]')}  |  ")
    escalation.add_run("Front office: ").bold = True
    escalation.add_run(f"{emergency.get('office_extension', '[blank]')}")

    doc.add_paragraph()

    doc.add_heading("Return Notes", level=2)
    return_notes = content.get("return_notes",
        "Please leave with the teacher: completed student work, attendance notes, any issues, and one sentence about how the lesson went.")
    doc.add_paragraph(return_notes)

    doc.add_paragraph()
    doc.add_paragraph(
        "Thank you for stepping in! Questions? See a nearby teacher or call the front office."
    )

    return doc


def _replace_starter_template_value(doc, label: str, value: str) -> None:
    """Populate the packaged starter template's label/value rows directly."""
    if not value:
        return
    target = label.strip().lower()
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) < 2:
                continue
            if row.cells[0].text.strip().lower() != target:
                continue
            row.cells[1].text = value
            return


def _replace_starter_template_return_notes(doc, return_notes: str) -> None:
    """Swap the starter template's baked-in return-notes paragraph when supplied."""
    if not return_notes:
        return
    for para in doc.paragraphs:
        if para.text.strip().startswith("Please leave a note for me"):
            para.text = return_notes
            return


def _build_from_template(fields: dict, content: dict, template_path: Path,
                        period: Optional[str] = None,
                        teacher_name: Optional[str] = None):
    """Fill template placeholders with content. Returns the Document."""
    doc = Document(str(template_path))

    learning_focus = (
        content.get("learning_focus") or fields.get("learning_intention", "")
    )
    activity_text = "\n".join(_activity_lines(fields, content))

    materials_list = content.get("materials") or fields.get("materials", [])
    materials_text = "\n".join(f"• {mat}" for mat in materials_list) if materials_list else ""
    emergency_text = _format_emergency_text(content)
    return_notes = content.get("return_notes", "")

    replacements = {
        "{{TEACHER}}": teacher_name or "",
        "{{SUBJECT}}": fields.get("subject", ""),
        "{{DATE}}": fields["date"],
        "{{PERIODS}}": content.get("periods") or period or "",
        "{{LEARNING_FOCUS}}": learning_focus,
        "{{ACTIVITY}}": activity_text,
        "{{MATERIALS}}": materials_text,
        "{{BACKUP}}": content.get("backup", ""),
        "{{EMERGENCY}}": emergency_text,
        "{{NEARBY_TEACHER}}": content.get("emergency", {}).get("nearby_teacher", ""),
        "{{OFFICE_EXTENSION}}": content.get("emergency", {}).get("office_extension", ""),
        "{{RETURN_NOTES}}": return_notes,
    }

    replace_docx_placeholders(doc, replacements)
    _replace_starter_template_value(doc, "Period(s)", content.get("periods") or period or "")
    _replace_starter_template_value(doc, "Today's activity", activity_text)
    _replace_starter_template_value(doc, "Materials", materials_text)
    _replace_starter_template_value(doc, "Backup activity", content.get("backup", ""))
    _replace_starter_template_value(doc, "If something breaks", emergency_text)
    _replace_starter_template_return_notes(doc, return_notes)

    return doc


def main() -> int:
    ap = argparse.ArgumentParser(description="Generate sub plan from lesson plan.")
    ap.add_argument("--plan", required=True, help="Path to markdown lesson plan")
    ap.add_argument("--date", required=True, help="Target date (YYYY-MM-DD)")
    ap.add_argument("--subject", help="Subject id (for display)")
    ap.add_argument("--output", required=True, help="Output .docx path")
    ap.add_argument("--template", help="Optional .docx template to use as base")
    ap.add_argument("--period", help="Period or time block")
    ap.add_argument("--teacher-name", help="Teacher name for header")
    ap.add_argument("--content-file", help="Path to Haiku-generated content JSON")
    ap.add_argument("--content-json", help="Inline Haiku-generated content JSON")
    ap.add_argument("--config", help="Config file for PII allowlist")
    ap.add_argument(
        "--allow-names",
        help=(
            "Comma-separated bare names to allow "
            "(teacher/public figures only; keyword-shaped PII still fails)"
        ),
    )
    ap.add_argument("--allow-anywhere", action="store_true", help="Allow writing outside ~/Documents/Lesson Plan Magic/outputs/")
    args = ap.parse_args()

    plan_path = Path(args.plan)
    try:
        md = load_plan_md(plan_path)
    except (FileNotFoundError, ValueError) as e:
        print(f"Error: {e}", file=sys.stderr)
        return 1

    try:
        fields = extract_day_fields(md, args.date)
    except ValueError as e:
        print(f"Error: {e}", file=sys.stderr)
        return 1
    if fields is None:
        print(f"Error: date {args.date} not found in plan.", file=sys.stderr)
        return 1

    if args.subject:
        fields["subject"] = args.subject

    try:
        content = load_and_validate_content(
            content_file=args.content_file,
            content_json=args.content_json,
            validator=validate_sub_plan_content,
        )
    except ContentValidationError as e:
        print(f"Error: invalid content JSON: {e}", file=sys.stderr)
        return 1

    try:
        cfg = load_config(args.config)
    except (FileNotFoundError, RuntimeError, ValueError) as e:
        print(f"Error: {e}", file=sys.stderr)
        return 1
    resolved_teacher, approved_names = resolve_allowlist(
        md, cfg, cli_allow_names=args.allow_names or "",
    )
    # --teacher-name CLI wins over config-derived teacher name.
    teacher_name = (args.teacher_name or resolved_teacher or "").strip()
    if args.teacher_name and resolved_teacher and args.teacher_name != resolved_teacher:
        approved_names.append(resolved_teacher)

    merged_text = md + "\n" + json.dumps(content, default=str)
    pii_error = scan_for_pii(merged_text, teacher_name=teacher_name, approved_names=approved_names)
    if pii_error:
        print(
            "Error: PII check failed. Remove student-identifying or sensitive information and retry.",
            file=sys.stderr,
        )
        return 1

    try:
        output_path = resolve_output_path(
            args.output, allow_anywhere=args.allow_anywhere
        )
    except ValueError as e:
        print(f"Error: {e}", file=sys.stderr)
        return 1

    mode_note = ""
    if args.template:
        tpl = Path(args.template)
        if tpl.exists():
            doc = _build_from_template(fields, content, tpl,
                                       period=args.period,
                                       teacher_name=teacher_name)
            mode_note = " (from template)"
        else:
            print(f"Warning: template not found ({tpl}), generating from scratch.", file=sys.stderr)
            doc = _build_from_scratch(fields, content,
                                      period=args.period,
                                      teacher_name=teacher_name)
    else:
        doc = _build_from_scratch(fields, content,
                                  period=args.period,
                                  teacher_name=teacher_name)

    unresolved = find_unresolved_docx_placeholders(doc)
    if unresolved:
        print(
            "Error: unresolved template placeholders remain after fill: "
            + ", ".join(unresolved[:10]),
            file=sys.stderr,
        )
        return 1

    pii_error = scan_docx_for_pii(
        doc, teacher_name=teacher_name, approved_names=approved_names,
    )
    if pii_error:
        print(
            "Error: PII check failed on rendered sub plan. Remove student-identifying or sensitive information and retry.",
            file=sys.stderr,
        )
        return 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Wrote{mode_note}: {output_path}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
