#!/usr/bin/env python3
"""parse_standards.py — PDF/DOCX/URL/text -> indexed standards JSON.

Usage:
  python parse_standards.py --input <path-or-url> [--subject <id>] [--cache-dir <dir>] [--pretty]
  cat text | python parse_standards.py --stdin [--subject <id>]

Internal cache format is DENSE (no indentation, short keys, nulls dropped) — the
plan pipeline reads it, not humans. Pass --pretty for readable stdout.

Dense schema:
  {s:{t,p?,u?,h}, f, g:[...], S:[{c, st, ft?, pg?}, ...]}
   s=source t=type p=path u=url h=hash
   f=framework g=grade_levels
   S=standards  c=code  st=short_text  ft=full_text  pg=source_page
"""
from __future__ import annotations
import argparse, hashlib, json, os, re, sys, tempfile
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Optional

_SHARED_DIR = Path(__file__).resolve().parents[3] / "shared"
if str(_SHARED_DIR) not in sys.path:
    sys.path.insert(0, str(_SHARED_DIR))

try:
    from runtime_bootstrap import ensure_plugin_runtime_or_exit
except ImportError:  # pragma: no cover - exercised by isolated script tests
    ensure_plugin_runtime_or_exit = None

if ensure_plugin_runtime_or_exit is not None:
    ensure_plugin_runtime_or_exit(__file__)

import safe_http

FRAMEWORKS = [
    ("NGSS", r"Next Generation Science Standards|NGSS"),
    ("Common Core ELA", r"Common Core State Standards.{0,40}English Language Arts|CCSS[.-]ELA"),
    ("Common Core Math", r"Common Core State Standards.{0,40}Math|CCSS[.-]Math"),
    ("Georgia Standards of Excellence", r"Georgia Standards of Excellence|GSE\b"),
    ("AP College Board", r"College Board|AP Course and Exam Description"),
    ("TEKS", r"Texas Essential Knowledge and Skills|\bTEKS\b"),
    ("C3 Social Studies", r"C3 Framework|College, Career, and Civic Life"),
]
CODE_PATTERNS = {
    "NGSS": r"\b((?:MS|HS|K|[1-5])-[A-Z]{2,4}\d*-\d+)\b",
    "Common Core ELA": r"\b(CCSS\.ELA-LITERACY\.[A-Z]{1,3}\.\d+\.\d+[a-z]?)\b",
    "Common Core Math": r"\b(CCSS\.MATH\.CONTENT\.[A-Z0-9.]+)\b",
    "Georgia Standards of Excellence": r"\b([A-Z]{2,8}\d*\.?[A-Z]?\.?\d+[a-z]?)\b",
    "AP College Board": r"\b(\d+\.\d+\.[A-Z]\.\d+)\b",
    "TEKS": r"\b\((\d{1,2}\)\([A-Z]\))\b",
    "C3 Social Studies": r"\b(D\d\.[A-Za-z]+\.\d+-\d+)\b",
}
GENERIC_CODE = r"^\s*([A-Z]{2,}[-.]?\d+[A-Z0-9.-]*)\s+(.{10,})"
MAX_FETCH_BYTES = 25 * 1024 * 1024
SUBJECT_ID_RE = re.compile(r"^[a-z0-9][a-z0-9-]*$")


@dataclass
class Standard:
    code: str
    short_text: str = ""
    full_text: str = ""
    grade_level: Optional[str] = None
    topic: Optional[str] = None
    sub_topic: Optional[str] = None
    performance_expectation: Optional[str] = None
    science_and_engineering_practices: Optional[list] = None
    crosscutting_concepts: Optional[list] = None
    source_page: Optional[int] = None


@dataclass
class Source:
    type: str
    path: Optional[str] = None
    url: Optional[str] = None
    hash: Optional[str] = None


@dataclass
class ParsedStandards:
    source: Source
    framework: str = "unknown"
    grade_levels: list = field(default_factory=list)
    standards: list = field(default_factory=list)


def detect_framework(text: str) -> str:
    head = text[:10000]
    for name, pat in FRAMEWORKS:
        if re.search(pat, head, re.IGNORECASE):
            return name
    return "unknown"


def detect_grade_levels(text: str) -> list:
    found = set()
    for m in re.finditer(r"\bGrade[s]?\s*(K|\d{1,2})(?:\s*[-\u2013]\s*(\d{1,2}))?\b", text):
        a, b = m.group(1), m.group(2)
        if a == "K":
            found.add("K")
            a = 0
        try:
            a = int(a); b = int(b) if b else a
            for g in range(a, b + 1):
                found.add(str(g))
        except (ValueError, TypeError):
            pass
    for m in re.finditer(r"\b(HS|MS|High School|Middle School|Elementary)\b", text, re.IGNORECASE):
        found.add({"hs": "9-12", "ms": "6-8", "high school": "9-12",
                   "middle school": "6-8", "elementary": "K-5"}[m.group(1).lower()])
    return sorted(found)


def extract_standards(text: str, framework: str, page_map: Optional[dict] = None) -> list:
    pat = CODE_PATTERNS.get(framework)
    out, seen = [], set()
    if pat:
        for m in re.finditer(pat, text):
            code = m.group(1)
            if code in seen:
                continue
            seen.add(code)
            start = max(0, m.start() - 10)
            end = min(len(text), m.start() + 400)
            chunk = text[start:end]
            para = re.split(r"\n{2,}", chunk, maxsplit=2)[0].strip()
            short = re.sub(r"\s+", " ", para[:200])
            full = re.sub(r"\s+", " ", para[:800])
            page = None
            if page_map:
                for p, pos in sorted(page_map.items()):
                    if pos <= m.start():
                        page = p
            out.append(Standard(code=code, short_text=short, full_text=full, source_page=page))
    else:
        for line in text.splitlines():
            m = re.match(GENERIC_CODE, line)
            if m:
                code = m.group(1)
                if code in seen:
                    continue
                seen.add(code)
                out.append(Standard(code=code, short_text=m.group(2)[:200],
                                    full_text=m.group(2)[:800]))
    return out


def _hash_bytes(b: bytes) -> str:
    return "sha256:" + hashlib.sha256(b).hexdigest()


def parse_pdf(path: Path) -> ParsedStandards:
    from pypdf import PdfReader
    reader = PdfReader(str(path))
    pages, page_map, cursor = [], {}, 0
    for i, page in enumerate(reader.pages, 1):
        t = page.extract_text() or ""
        page_map[i] = cursor
        pages.append(t)
        cursor += len(t) + 2
    text = "\n\n".join(pages)
    fw = detect_framework(text)
    return ParsedStandards(
        source=Source(type="pdf", path=str(path), hash=_hash_bytes(path.read_bytes())),
        framework=fw, grade_levels=detect_grade_levels(text),
        standards=extract_standards(text, fw, page_map))


def parse_docx(path: Path) -> ParsedStandards:
    from docx import Document
    doc = Document(str(path))
    paras = [p.text for p in doc.paragraphs if p.text.strip()]
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paras.append(cell.text)
    text = "\n\n".join(paras)
    fw = detect_framework(text)
    return ParsedStandards(
        source=Source(type="docx", path=str(path), hash=_hash_bytes(path.read_bytes())),
        framework=fw, grade_levels=detect_grade_levels(text),
        standards=extract_standards(text, fw))


def _is_private_ip(hostname: str) -> bool:
    """Resolve hostname and reject private/loopback/link-local addresses."""
    return safe_http.is_private_ip(hostname)


def _validate_url(url: str) -> safe_http.ValidatedURL:
    """Block non-http(s) schemes and return a pinned public target."""
    return safe_http.validate_url(url)


def parse_url(url: str, cache_dir: Optional[Path] = None) -> ParsedStandards:
    current_target = _validate_url(url)
    current_url = current_target.url
    status, headers, body = safe_http.fetch_url(
        current_target,
        timeout=15,
        user_agent="LessonPlanMagic/0.2",
        max_body_bytes=MAX_FETCH_BYTES,
    )
    # Follow redirects manually, revalidating each hop
    hops = 0
    while status in safe_http.REDIRECT_STATUSES and hops < 5:
        next_url = headers.get("Location", "")
        if not next_url:
            break
        # Resolve relative redirects
        from urllib.parse import urljoin
        next_url = urljoin(current_url, next_url)
        current_target = _validate_url(next_url)
        status, headers, body = safe_http.fetch_url(
            current_target,
            timeout=15,
            user_agent="LessonPlanMagic/0.2",
            max_body_bytes=MAX_FETCH_BYTES,
        )
        current_url = current_target.url
        hops += 1
    if status >= 400:
        raise RuntimeError(f"HTTP {status}")
    ctype = headers.get("Content-Type", "").lower()
    suffix = ".pdf" if "pdf" in ctype or current_url.lower().endswith(".pdf") else \
             ".docx" if "wordprocessingml" in ctype or current_url.lower().endswith(".docx") else \
             ".html"
    fd, tmp_name = tempfile.mkstemp(suffix=suffix)
    os.close(fd)
    tmp = Path(tmp_name)
    tmp.write_bytes(body)
    try:
        if suffix == ".pdf":
            ps = parse_pdf(tmp)
        elif suffix == ".docx":
            ps = parse_docx(tmp)
        else:
            try:
                from bs4 import BeautifulSoup
                text = BeautifulSoup(body, "html.parser").get_text("\n")
            except Exception:
                text = body.decode("utf-8", errors="replace")
            fw = detect_framework(text)
            ps = ParsedStandards(
                source=Source(type="url", url=current_url, hash=_hash_bytes(body)),
                framework=fw, grade_levels=detect_grade_levels(text),
                standards=extract_standards(text, fw))
        ps.source = Source(type="url", url=current_url, hash=_hash_bytes(body))
        return ps
    finally:
        try:
            tmp.unlink()
        except OSError:
            pass


def parse_text(text: str) -> ParsedStandards:
    fw = detect_framework(text)
    return ParsedStandards(
        source=Source(type="text", hash=_hash_bytes(text.encode())),
        framework=fw, grade_levels=detect_grade_levels(text),
        standards=extract_standards(text, fw))


def _to_dict(ps: ParsedStandards) -> dict:
    """Full/pretty dict — long keys, all fields."""
    return {"source": asdict(ps.source), "framework": ps.framework,
            "grade_levels": ps.grade_levels,
            "standards": [asdict(s) for s in ps.standards]}


def _to_dense(ps: ParsedStandards) -> dict:
    """Dense dict for internal cache — short keys, nulls dropped."""
    src = {"t": ps.source.type}
    if ps.source.path: src["p"] = ps.source.path
    if ps.source.url:  src["u"] = ps.source.url
    if ps.source.hash: src["h"] = ps.source.hash
    stds = []
    for s in ps.standards:
        d = {"c": s.code, "st": s.short_text}
        if s.full_text and s.full_text != s.short_text: d["ft"] = s.full_text
        if s.source_page is not None: d["pg"] = s.source_page
        stds.append(d)
    return {"s": src, "f": ps.framework, "g": ps.grade_levels, "S": stds}


def validate_subject_id(subject_id: str) -> str:
    if not SUBJECT_ID_RE.fullmatch(subject_id or ""):
        raise ValueError(
            "subject id must be kebab-case (lowercase letters, digits, hyphens only)"
        )
    return subject_id


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--input", help="Path or URL")
    ap.add_argument("--stdin", action="store_true")
    ap.add_argument("--subject", help="subject id for cache filename")
    ap.add_argument("--cache-dir", help="cache directory")
    ap.add_argument("--pretty", action="store_true",
                    help="Pretty-print stdout (default: dense); cache is always dense")
    a = ap.parse_args()
    try:
        subject_id = validate_subject_id(a.subject) if a.subject else None
    except ValueError as exc:
        print(f"error: {exc}", file=sys.stderr)
        return 1
    if a.stdin:
        ps = parse_text(sys.stdin.read())
    elif a.input:
        if a.input.startswith(("http://", "https://")):
            ps = parse_url(a.input, Path(a.cache_dir) if a.cache_dir else None)
        else:
            p = Path(a.input)
            if not p.exists():
                print(f"error: file not found: {p}", file=sys.stderr)
                return 1
            ext = p.suffix.lower()
            if ext == ".pdf":
                ps = parse_pdf(p)
            elif ext == ".docx":
                ps = parse_docx(p)
            elif ext in (".txt", ".md"):
                ps = parse_text(p.read_text())
            else:
                print(f"error: unsupported extension: {ext}", file=sys.stderr)
                return 1
    else:
        ap.error("--input or --stdin required")
        return 1
    dense = _to_dense(ps)
    if a.pretty:
        print(json.dumps(_to_dict(ps), indent=2))
    else:
        print(json.dumps(dense, separators=(",", ":")))
    if a.cache_dir and subject_id:
        out = Path(a.cache_dir)
        out.mkdir(parents=True, exist_ok=True)
        (out / f"{subject_id}.parsed.json").write_text(
            json.dumps(dense, separators=(",", ":"))
        )
    return 0


if __name__ == "__main__":
    sys.exit(main())
