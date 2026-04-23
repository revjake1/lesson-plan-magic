#!/usr/bin/env python3
"""
Ingest teacher's past .docx lesson plans and extract voice profile dimensions.
Outputs human-editable voice-profile.md and a dense voice-profile.json sidecar
for use by the lesson-planner skill.
"""

from __future__ import annotations

import argparse
import json
import logging
import re
import sys
import zipfile
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path
from statistics import median
from typing import Dict, List, Tuple

_SHARED_DIR = Path(__file__).resolve().parents[3] / "shared"
if str(_SHARED_DIR) not in sys.path:
    sys.path.insert(0, str(_SHARED_DIR))

try:
    from runtime_bootstrap import ensure_plugin_runtime_or_exit
except ImportError:  # pragma: no cover - exercised by isolated script tests
    ensure_plugin_runtime_or_exit = None

if ensure_plugin_runtime_or_exit is not None:
    ensure_plugin_runtime_or_exit(__file__)

from defusedxml import ElementTree as ET
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table

# Hardcoded activity keywords and redaction allowlist
ACTIVITY_KEYWORDS = [
    "think-pair-share", "think pair share", "turn and talk", "pair-share",
    "gallery walk", "jigsaw", "fishbowl", "socratic seminar", "four corners",
    "exit ticket", "exit slip", "do-now", "do now", "bell ringer", "bellringer",
    "warm-up", "warm up", "warmup",
    "3-2-1", "quick write", "carousel", "stations",
    "entry ticket", "anchor chart", "word wall", "frayer",
    "claim evidence reasoning", "cer", "rally robin", "numbered heads",
    "close reading", "annotation", "notice and wonder",
    "chalk talk", "see think wonder", "i notice i wonder",
    "sketchnote", "one pager", "gallery",
]

DEFAULT_NAME_ALLOWLIST = {
    "Rosa Parks", "Abraham Lincoln", "Martin Luther", "William Shakespeare",
    "Mark Twain", "Langston Hughes", "Maya Angelou", "Frederick Douglass",
    "Harriet Tubman",
}

# Case-insensitive structural vocabulary. If either word of a Cap+Cap name
# candidate is in this set, the full-name pass suppresses the redaction.
# Prevents template headings like ``LEARNING INTENTION`` / ``SUCCESS
# CRITERIA`` / ``EXIT TICKET`` (or title-case equivalents) from being
# mis-redacted as student names.
_STRUCTURAL_WORDS_LOWER = {
    # Days & months
    "monday", "tuesday", "wednesday", "thursday", "friday",
    "saturday", "sunday",
    "january", "february", "march", "april", "may", "june",
    "july", "august", "september", "october", "november", "december",
    # Plan field labels
    "subject", "teacher", "date", "week", "day", "period", "periods",
    "standards", "learning", "intention", "intentions",
    "success", "criteria", "agenda", "materials", "differentiation",
    "evidence", "assessment", "unit", "lesson", "bell", "ringer",
    "warm", "warmup", "exit", "ticket", "homework",
    "objective", "objectives", "goal", "goals",
    "framework", "frameworks", "schedule", "room",
    "chapter", "section", "module", "topic", "activity", "activities",
    "do", "now", "opener", "closer", "main", "intro", "outro",
    "reflection", "reflections", "notes", "note", "of",
    # Education terms
    "grade", "level", "course", "semester", "quarter", "trimester",
    "quiz", "exam", "project", "group", "groups", "pair", "pairs",
    # Compliance shorthand
    "iep", "504", "ell", "sped", "bip", "mtss", "rti", "ese", "gt",
    # Subjects
    "math", "science", "english", "history", "social", "studies",
    "chemistry", "biology", "physics", "algebra", "geometry", "calculus",
    "literature", "writing", "reading", "art", "music", "pe",
    "health", "economics", "government", "world", "american", "us",
    # Sentence-starter verbs / transitions common in plan text
    "review", "introduce", "discuss", "explore", "explain", "present",
    "summarize", "analyze", "compare", "contrast", "evaluate", "apply",
    "create", "describe", "identify", "show", "today", "tomorrow",
    "yesterday", "consider", "remember", "notice", "think", "write",
    "read", "watch", "listen", "study", "complete", "begin", "start",
    "finish", "close", "open", "ask", "answer", "choose", "see",
    "visit", "go", "work", "continue", "finally", "first", "second",
    "third", "next", "then", "after", "before", "during", "while",
    "when", "where", "why", "how", "what", "who", "which",
    "let", "have", "give", "take", "make", "use", "pick",
    "team", "class", "include", "exclude", "provide", "allow",
    "encourage", "model", "demonstrate", "practice", "extend", "reduce",
    "spiral", "reteach", "test",
    # Student-keyword tokens themselves (prevents "STUDENT WORK" being
    # misread as a two-name pair after the keyword-anchored pass runs).
    "student", "students", "pupil", "pupils", "kid", "kids",
    "child", "children", "learner", "learners",
}

logging.basicConfig(level=logging.WARNING, format="%(levelname)s: %(message)s")
logger = logging.getLogger(__name__)

_DOCX_CORE_NS = "http://schemas.openxmlformats.org/package/2006/metadata/core-properties"
_DOCX_DC_NS = "http://purl.org/dc/elements/1.1/"
_DOCX_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
MAX_DOCX_XML_ENTRY_BYTES = 2 * 1024 * 1024
MAX_DOCX_TOTAL_UNCOMPRESSED_BYTES = 40 * 1024 * 1024
MAX_DOCX_ENTRY_COUNT = 2048


def _print_status(message: str) -> None:
    """Emit a console status line using ASCII-only text."""
    print(message.encode("ascii", "backslashreplace").decode("ascii"))


def _validate_docx_package(doc_path: Path) -> None:
    with zipfile.ZipFile(doc_path) as zf:
        infos = zf.infolist()
        if len(infos) > MAX_DOCX_ENTRY_COUNT:
            raise ValueError(
                f"{doc_path.name} contains too many zip entries ({len(infos)})."
            )
        total_bytes = sum(info.file_size for info in infos)
        if total_bytes > MAX_DOCX_TOTAL_UNCOMPRESSED_BYTES:
            raise ValueError(
                f"{doc_path.name} expands past the {MAX_DOCX_TOTAL_UNCOMPRESSED_BYTES}-byte safety limit."
            )
        for info in infos:
            if info.filename.endswith(".xml") and info.file_size > MAX_DOCX_XML_ENTRY_BYTES:
                raise ValueError(
                    f"{doc_path.name} contains oversized XML part {info.filename}."
                )


def _read_capped_zip_entry(
    zf: zipfile.ZipFile, name: str, *, max_bytes: int = MAX_DOCX_XML_ENTRY_BYTES
) -> bytes:
    info = zf.getinfo(name)
    if info.file_size > max_bytes:
        raise ValueError(f"Oversized XML part: {name}")
    with zf.open(info) as fh:
        data = fh.read(max_bytes + 1)
    if len(data) > max_bytes:
        raise ValueError(f"Oversized XML part: {name}")
    return data


def extract_docx_metadata_authors(doc_path: Path) -> List[str]:
    """Return distinct author-like metadata values from a .docx package.

    Scans core properties plus WordprocessingML parts for ``w:author``
    attributes so comment authors and tracked-change authors get fed
    through the redaction pass without skewing the visible-text analysis.
    """
    try:
        _validate_docx_package(doc_path)
        with zipfile.ZipFile(doc_path) as zf:
            values: list[str] = []

            try:
                core_root = ET.fromstring(_read_capped_zip_entry(zf, "docProps/core.xml"))
            except (KeyError, ET.ParseError, ValueError):
                core_root = None
            if core_root is not None:
                for tag in (
                    f"{{{_DOCX_DC_NS}}}creator",
                    f"{{{_DOCX_CORE_NS}}}lastModifiedBy",
                ):
                    for node in core_root.findall(f".//{tag}"):
                        text = (node.text or "").strip()
                        if text:
                            values.append(text)

            for name in zf.namelist():
                if not (name.startswith("word/") and name.endswith(".xml")):
                    continue
                try:
                    root = ET.fromstring(_read_capped_zip_entry(zf, name))
                except (ET.ParseError, ValueError):
                    continue
                for node in root.iter():
                    author = (
                        node.attrib.get(f"{{{_DOCX_W_NS}}}author")
                        or node.attrib.get("author")
                        or ""
                    ).strip()
                    if author:
                        values.append(author)
    except (OSError, zipfile.BadZipFile, ValueError):
        return []

    deduped: list[str] = []
    seen: set[str] = set()
    for value in values:
        if value not in seen:
            deduped.append(value)
            seen.add(value)
    return deduped


def extract_text_from_docx(doc_path: Path) -> Tuple[str, int, int]:
    """Extract visible text from paragraphs and tables.

    Metadata authors (core properties, comments, tracked changes) are
    scanned separately via ``extract_docx_metadata_authors`` so they can
    be redacted without polluting the voice-profile analysis.
    """
    try:
        _validate_docx_package(doc_path)
        doc = Document(doc_path)
    except Exception as e:
        logger.warning(f"Failed to open {doc_path.name}: {e}")
        return "", 0, 0

    text_parts = []
    table_count = 0
    total_rows = 0

    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    for table in doc.tables:
        table_count += 1
        total_rows += len(table.rows)
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text)

    return "\n".join(text_parts), table_count, total_rows


def redact_names(text: str, redact: bool = True, teacher_name: str = None) -> Tuple[str, int]:
    """Redact plausible student PII. Return (redacted_text, count).

    Covers: keyword-anchored roster lines ("Student: John Doe" /
    "STUDENT: JOHN DOE"), Firstname Lastname (title-case OR all-caps),
    lone first names in student-context lines, initials (A.B., A.B.C.),
    email addresses, numeric student IDs, and accommodation/IEP/504
    strings that imply disability information.
    """
    if not redact:
        return text, 0

    allowlist = set(DEFAULT_NAME_ALLOWLIST)
    # Build a set of allowlisted first names so we don't redact "Rosa" when
    # it appears solo in a line that also mentions a historical figure.
    allowlist_firsts = {n.split()[0] for n in allowlist if " " in n}
    if teacher_name:
        allowlist.add(teacher_name)
        allowlist_firsts.add(teacher_name.split()[0])

    # Case-insensitive mirrors so all-caps variants ("ROSA PARKS",
    # "JANE DOE" when "Jane Doe" is the teacher) compare correctly.
    allowlist_lower = {n.lower() for n in allowlist}
    allowlist_firsts_lower = {f.lower() for f in allowlist_firsts}

    redaction_count = 0

    # 0. Keyword-anchored roster — "Student: John Doe", "STUDENT: JOHN DOE",
    #    "learner: smith". The keyword itself denotes student context, so
    #    this runs BEFORE the allowlist / structural filters and redacts
    #    the whole keyword+name unit unconditionally. Matches 1-2
    #    capitalized name tokens after the keyword (covers both title-case
    #    and all-caps via [A-Z][A-Za-z'\-]+).
    def _replace_keyword_roster(match):
        nonlocal redaction_count
        redaction_count += 1
        return "[STUDENT]"

    result = re.sub(
        r"\b(?i:student|pupil|kid|child|learner)s?[:;][ \t]*"
        r"[A-Z][A-Za-z'\-]+(?:[ \t]+[A-Z][A-Za-z'\-]+)?",
        _replace_keyword_roster, text,
    )

    # 1. Full "Firstname Lastname" — title-case AND all-caps. Structural
    #    vocabulary (case-insensitive) and allowlist (case-insensitive)
    #    suppress matches that are actually section headings or historical
    #    figures. Broadened from the original [A-Z][a-z]+ pattern so
    #    "JOHN DOE" in uppercase plans can no longer slip through.
    def _replace_full(match):
        nonlocal redaction_count
        full_name = match.group(1)
        parts = full_name.split()
        if any(p.lower() in _STRUCTURAL_WORDS_LOWER for p in parts):
            return full_name
        if full_name.lower() in allowlist_lower:
            return full_name
        redaction_count += 1
        return "[STUDENT]"

    result = re.sub(
        r"\b([A-Z][A-Za-z'\-]+\s+[A-Z][A-Za-z'\-]+)\b",
        _replace_full, result,
    )

    # 2. Initials like "A.B.", "A.B.C.", "J. S." (common in IEP references)
    def _replace_initials(match):
        nonlocal redaction_count
        redaction_count += 1
        return "[STUDENT]"

    result = re.sub(r"\b([A-Z]\.)\s*([A-Z]\.)\s*([A-Z]\.)?\b", _replace_initials, result)

    # 3. Email addresses (student emails often contain names)
    def _replace_email(match):
        nonlocal redaction_count
        redaction_count += 1
        return "[EMAIL]"

    result = re.sub(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b", _replace_email, result)

    # 4. Numeric student IDs (6-10 digit numbers, optionally prefixed by
    #    "ID", "Student #", etc.)
    def _replace_id(match):
        nonlocal redaction_count
        redaction_count += 1
        return "[STUDENT-ID]"

    result = re.sub(
        r"(?:(?:student|ID|id)\s*#?\s*)(\d{6,10})\b",
        _replace_id, result
    )
    # Standalone long digit strings that look like IDs (not years, not page numbers)
    result = re.sub(r"\b(?<!\d)(\d{7,10})(?!\d)\b", _replace_id, result)

    # 5. Accommodation / IEP / 504 strings that describe specific student
    #    disabilities or services — these are FERPA-protected when linked to
    #    an identifiable student, even after the name is gone, because the
    #    combination of class + accommodation can re-identify.
    accommodation_pattern = (
        r"(?i)\b(?:IEP|504 plan|accommodation|modification|"
        r"speech therapy|occupational therapy|OT services|"
        r"extended time|preferential seating|"
        r"behavior intervention plan|BIP|"
        r"functional behavior assessment|FBA)\b"
        r"[^.\n]{0,120}"
    )
    def _replace_accommodation(match):
        nonlocal redaction_count
        redaction_count += 1
        return "[ACCOMMODATION REDACTED]"

    result = re.sub(accommodation_pattern, _replace_accommodation, result)

    # 6. Lone first names on lines that look student-specific (attendance,
    #    grouping, call-on lists). Heuristic: a capitalized word preceded by
    #    a bullet/number and not in the historical-figure allowlist. Matches
    #    title-case "Marcus" and all-caps "MARCUS"; the all-caps branch
    #    additionally suppresses structural vocabulary so heading-style
    #    bullets like "- TODAY" or "- EXIT" aren't mis-redacted.
    def _replace_lone_name(match):
        nonlocal redaction_count
        name = match.group(2)
        if name.lower() in allowlist_firsts_lower:
            return match.group(0)
        if name.isupper() and name.lower() in _STRUCTURAL_WORDS_LOWER:
            return match.group(0)
        redaction_count += 1
        return match.group(1) + "[STUDENT]"

    # `[-*•]\s*` — bullet markers need optional whitespace before the name
    # (the numbered-list branch already includes `\s*`). Without `\s*` here,
    # the canonical form "- Marcus" leaks through unredacted.
    result = re.sub(
        r"(^\s*(?:[-*•]\s*|\d+[.)]\s*))([A-Z][a-z]{2,}|[A-Z]{2,})\b",
        _replace_lone_name, result, flags=re.MULTILINE
    )

    return result, redaction_count


def extract_layout_signature(docs_data: List[Tuple[str, int, int]]) -> Dict:
    """Analyze table density, section order, bullet ratio, heading style."""
    tables = [d[1] for d in docs_data]
    rows = [d[2] for d in docs_data]

    avg_tables = sum(tables) / len(tables) if tables else 0
    avg_rows = median(rows) if rows else 0

    all_text = "\n".join([d[0] for d in docs_data])
    # Count lines that START with a bullet/number-dot (not anywhere).
    layout_bullet_pattern = r"^\s*(?:[-*•]|\d+\.)\s+"
    bullet_lines = len(re.findall(layout_bullet_pattern, all_text, re.MULTILINE))
    prose_lines = len([
        l for l in all_text.split("\n")
        if l.strip() and not re.match(layout_bullet_pattern, l)
    ])
    bullet_ratio = (
        bullet_lines / (bullet_lines + prose_lines)
        if (bullet_lines + prose_lines) > 0
        else 0
    )

    # Simple heading detection: lines containing "Heading" style, bold markers, or all-caps
    heading_markers = len(re.findall(r"^\s*[A-Z\s]{5,}$", all_text, re.MULTILINE))

    return {
        "avg_tables_per_doc": round(avg_tables, 2),
        "avg_rows_per_table": round(avg_rows, 2),
        "bullet_ratio": round(bullet_ratio, 2),
        "heading_markers_count": heading_markers,
    }


def extract_voice_signature(docs_data: List[Tuple[str, int, int]]) -> Dict:
    """Analyze script vs outline, contractions, second-person, signature phrases, warmth."""
    all_text = "\n".join([d[0] for d in docs_data])
    word_count = len(all_text.split())

    # Script vs outline: median words per bullet
    # Match dash/asterisk/bullet OR numbered items. The previous regex
    # failed on numbered lists entirely (see audit D4).
    bullet_pattern = r"^\s*(?:[-*•]|\d+\.)\s+(.+)"
    bullets = re.findall(bullet_pattern, all_text, re.MULTILINE)
    median_words_per_bullet = median([len(b.split()) for b in bullets]) if bullets else 0

    # Contraction rate
    contractions = r"\b(don't|won't|they'll|can't|that's|it's|I'll|we'll|you're|there's)\b"
    contraction_count = len(re.findall(contractions, all_text, re.IGNORECASE))
    contraction_rate = (contraction_count / word_count * 1000) if word_count > 0 else 0

    # Second-person rate
    second_person = r"\b(you|your|students will|students do)\b"
    second_person_count = len(re.findall(second_person, all_text, re.IGNORECASE))
    second_person_rate = (second_person_count / word_count * 1000) if word_count > 0 else 0

    # Signature phrases (trigrams, ≥2 plans)
    stopwords = {"the", "a", "an", "of", "and", "or", "to", "in", "on", "for", "at", "is", "are", "be", "this", "that", "with", "by"}
    tokens = re.sub(r"[^\w\s']", " ", all_text.lower()).split()
    tokens = [t for t in tokens if len(t) > 1]
    trigrams = [tuple(tokens[i:i+3]) for i in range(len(tokens) - 2)]
    trigrams = [t for t in trigrams if not all(tok in stopwords for tok in t)]
    trigram_counts = Counter(trigrams)
    top_trigrams = [" ".join(t) for t, c in trigram_counts.most_common(10) if c >= 2]

    # Warmth markers
    #
    # NB: em-dashes are NOT counted as warmth. Em-dash overuse is one of
    # the classic AI writing tells — counting them would amplify exactly
    # the pattern the downstream pipeline is supposed to avoid. See
    # audit D3 + the humanizer skill.
    parenthetical_pattern = (
        r"\([^)]*(?:bonus|groan|ha|haha|cheers|yay|laugh|sigh|oof|nice)\)"
    )
    warmth_parens = len(re.findall(parenthetical_pattern, all_text, re.IGNORECASE))
    exclamation_marks = len(re.findall(r"!", all_text))
    # Sincere-warmth inline markers: "ha,", "(bonus)", etc. — already
    # covered in the parenthetical pattern for the paren form. Also
    # count bare exclamations sparingly (they count a third).
    warmth_count = warmth_parens + (exclamation_marks // 3)
    warmth_rate = (warmth_count / word_count * 1000) if word_count > 0 else 0

    return {
        "script_vs_outline_median_words": round(median_words_per_bullet, 1),
        "contraction_rate_per_1k": round(contraction_rate, 2),
        "second_person_rate_per_1k": round(second_person_rate, 2),
        "signature_phrases": top_trigrams,
        "warmth_markers_per_1k": round(warmth_rate, 2),
    }


def extract_activity_library(docs_data: List[Tuple[str, int, int]], source_files: List[str]) -> Dict:
    """Identify recurring activities across plans."""
    activity_counts = Counter()
    activity_sources = {kw: [] for kw in ACTIVITY_KEYWORDS}

    for (text, _, _), filename in zip(docs_data, source_files):
        text_lower = text.lower()
        for activity in ACTIVITY_KEYWORDS:
            if activity in text_lower:
                activity_counts[activity] += 1
                activity_sources[activity].append(filename)

    # Return only activities appearing ≥2 times
    recurring = {
        activity: {
            "count": activity_counts[activity],
            "sources": list(set(activity_sources[activity])),
        }
        for activity in ACTIVITY_KEYWORDS
        if activity_counts[activity] >= 2
    }
    return recurring


def extract_pacing(docs_data: List[Tuple[str, int, int]], source_files: List[str]) -> Dict:
    """Infer pacing from dates and unit/topic mentions."""
    date_patterns = [
        r"(\d{4})-(\d{1,2})-(\d{1,2})",  # YYYY-MM-DD
        r"(\d{1,2})-(\d{1,2})-(\d{2})",  # M-D-YY
        r"Week of\s+(\w+\s+\d{1,2})",    # Week of Month Day
    ]

    pacing = []
    skipped = 0

    for text, filename in zip([d[0] for d in docs_data], source_files):
        date_match = None
        for pattern in date_patterns:
            match = re.search(pattern, text)
            if match:
                date_match = match.group(0)
                break

        if not date_match:
            skipped += 1
            continue

        # Extract first topic/unit from first 200 chars
        first_lines = text[:500].split("\n")
        topic = next((line.strip() for line in first_lines if line.strip() and len(line.strip()) > 5), "Unknown")

        pacing.append({"date": date_match, "topic": topic})

    pacing.sort(key=lambda x: x["date"])
    return {"entries": pacing, "skipped": skipped}


def format_markdown_output(
    layout: Dict, voice: Dict, activities: Dict, pacing: Dict,
    source_count: int, source_files: List[str], redactions: int,
) -> str:
    """Format voice profile as markdown."""
    yaml_files = ", ".join([f'"{f}"' for f in source_files[:10]])
    if len(source_files) > 10:
        yaml_files += f", ... ({len(source_files) - 10} more)"

    md = f"""---
generated_at: {datetime.now(timezone.utc).isoformat().replace("+00:00", "Z")}
source_count: {source_count}
source_files: [{yaml_files}]
redactions: {redactions}
dimensions: [layout, voice, activities, pacing]
---

# Voice Profile - extracted from {source_count} past plans

> This file is human-editable. The lesson-planner skill loads it as context before generating new plans. Edit freely; your changes override the auto-extracted values on the next run.

## Voice signature

- **Script vs. outline**: median {voice['script_vs_outline_median_words']} words per bullet
- **Register**: {voice['contraction_rate_per_1k']:.2f} contractions/1k words; {voice['second_person_rate_per_1k']:.0f} you/your-rate/1k
- **Warmth**: {voice['warmth_markers_per_1k']:.2f} warmth markers/1k
- **Signature phrases** (top):
"""

    for phrase in voice["signature_phrases"][:5]:
        md += f"  - {phrase}\n"

    md += f"""
## Layout signature

- **Avg tables per doc**: {layout['avg_tables_per_doc']}
- **Avg rows per table**: {layout['avg_rows_per_table']}
- **Bullet ratio**: {layout['bullet_ratio']}
- **Heading markers found**: {layout['heading_markers_count']}

## Activity library

| Activity | Count | Sources |
|---|---|---|
"""

    for activity, data in sorted(activities.items(), key=lambda x: x[1]["count"], reverse=True)[:10]:
        sources_str = ", ".join(data["sources"][:3])
        if len(data["sources"]) > 3:
            sources_str += f" (+{len(data['sources']) - 3})"
        md += f"| {activity} | {data['count']} | {sources_str} |\n"

    md += "\n## Pacing (inferred)\n\n"
    for entry in pacing["entries"][:10]:
        md += f"- {entry['date']} -> {entry['topic']}\n"

    if pacing["skipped"] > 0:
        md += f"\n*Note: {pacing['skipped']} plan(s) had no parseable date.*\n"

    md += f"""
## Notes

- Redacted {redactions} plausible student names before analysis.
- Source sample size: {source_count} plans -- """

    if source_count < 3:
        md += f"**Warning: Only {source_count} plans analyzed. Voice match will be approximate.**"
    else:
        md += "reliable match expected."

    md += "\n"
    return md


def main():
    parser = argparse.ArgumentParser(description="Ingest past lesson plans and extract voice profile.")
    parser.add_argument("--input-dir", required=True, help="Directory containing .docx files")
    parser.add_argument("--output", required=True, help="Output path for voice-profile.md")
    parser.add_argument("--min-plans", type=int, default=3, help="Warn if fewer plans found (default: 3)")
    parser.add_argument("--redact-names", action="store_true", default=True, help="Redact student names (default: True)")
    parser.add_argument("--no-redact-names", dest="redact_names", action="store_false", help="Skip name redaction")
    parser.set_defaults(json_sidecar=True)
    parser.add_argument(
        "--json-sidecar",
        dest="json_sidecar",
        action="store_true",
        help="Write voice-profile.json (default: on).",
    )
    parser.add_argument(
        "--no-json-sidecar",
        dest="json_sidecar",
        action="store_false",
        help="Skip writing voice-profile.json.",
    )
    parser.add_argument("--teacher-name", help="Teacher's name to allowlist during redaction")

    args = parser.parse_args()

    input_dir = Path(args.input_dir)
    output_path = Path(args.output)

    if not input_dir.is_dir():
        logger.error(f"Input directory not found: {input_dir}")
        return 1

    docx_files = sorted(input_dir.glob("*.docx"))
    if not docx_files:
        logger.error(f"No .docx files found in {input_dir}")
        return 1

    # Extract and redact
    docs_data = []
    source_files = []
    total_redactions = 0

    for docx_file in docx_files:
        text, table_count, total_rows = extract_text_from_docx(docx_file)
        metadata_text = "\n".join(extract_docx_metadata_authors(docx_file))
        _, metadata_redactions = redact_names(
            metadata_text, args.redact_names, args.teacher_name
        )
        total_redactions += metadata_redactions
        if not text:
            continue

        redacted_text, redactions = redact_names(text, args.redact_names, args.teacher_name)
        total_redactions += redactions
        docs_data.append((redacted_text, table_count, total_rows))
        source_files.append(docx_file.name)

    if not docs_data:
        logger.error("No readable documents extracted.")
        return 1

    if len(docs_data) < args.min_plans:
        logger.warning(f"Only {len(docs_data)} plans found; {args.min_plans} recommended for reliable profile.")

    # Extract dimensions
    layout = extract_layout_signature(docs_data)
    voice = extract_voice_signature(docs_data)
    activities = extract_activity_library(docs_data, source_files)
    pacing = extract_pacing(docs_data, source_files)

    # Generate markdown
    md_content = format_markdown_output(layout, voice, activities, pacing, len(docs_data), source_files, total_redactions)

    # Write outputs
    try:
        output_path.write_text(md_content, encoding="utf-8")
        _print_status(f"OK: Voice profile written to {output_path}")
    except Exception as e:
        logger.error(f"Failed to write {output_path}: {e}")
        return 1

    if args.json_sidecar:
        # Dense JSON sidecar for internal plan-pipeline use (the .md is for humans).
        # Short keys; drop empty collections; no indentation.
        json_path = output_path.with_suffix(".json")
        def _nz(v):
            return v not in (None, "", [], {}, 0)
        sidecar = {
            "ts": datetime.now(timezone.utc).isoformat().replace("+00:00", "Z"),
            "n": len(docs_data),
            "L": layout,
            "V": voice,
            "A": {k: v["count"] for k, v in activities.items()},
            "P": [{"d": e["date"], "t": e["topic"]} for e in pacing["entries"]],
        }
        if total_redactions:
            sidecar["r"] = total_redactions
        sidecar = {k: v for k, v in sidecar.items() if _nz(v)}
        try:
            json_path.write_text(
                json.dumps(sidecar, separators=(",", ":"), sort_keys=True),
                encoding="utf-8",
            )
            _print_status(f"OK: JSON sidecar written to {json_path}")
        except Exception as e:
            logger.error(f"Failed to write {json_path}: {e}")
            return 1

    _print_status(
        f"OK: Analyzed {len(docs_data)} plans, redacted {total_redactions} names."
    )
    return 0


if __name__ == "__main__":
    exit(main())
