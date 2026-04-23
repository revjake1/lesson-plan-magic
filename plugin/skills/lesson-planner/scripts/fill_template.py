"""
Fill a district .docx lesson-plan template from markdown lesson plans.

Supports:
  * explicit {{PLACEHOLDERS}} (preferred)
  * header-matching auto-map in tables (per-table OR per-row day cardinality)
Enforces hard PII scan before writing output.

Does NOT silently fall through to append-mode on mapping failure — that
violates the plugin's hard rule #3. Instead, reports an explicit error
listing unmatched day/section pairs so the teacher can fix the template
or the mapping.

Optional post-write config mutations (opt-in via flags):
  --mark-verified <subject-id>
      Set subjects[id=<subject-id>].template.mapping_verified: true
      in the teacher's config.yaml.
  --update-pacing <subject-id>
      Set subjects[id=<subject-id>].pacing.last_planned_week_end to the
      ISO date of the last day in the plan.
  --config PATH
      Override the default ~/Documents/Lesson Plan Magic/config.yaml.
  --allow-anywhere
      Explicitly allow writing the output .docx and .plan.md/.plan.json sidecars
      outside ~/Documents/Lesson Plan Magic/outputs/.
  --skip-docx-scan
      Skip scanning the template's own existing text before filling it.
      Intended only for debugging a known-clean legacy template.
  --no-sidecar
      Opt out of writing plan sidecars .plan.md and .plan.json
      (default: sidecars ON). These files are a private bridge to the
      classroom-artifacts skill.

Python: requires >=3.9.
"""

from __future__ import annotations

import argparse
import copy
import json
import os
import re
import sys
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable

_SHARED_DIR = Path(__file__).resolve().parents[3] / "shared"
if str(_SHARED_DIR) not in sys.path:
    sys.path.insert(0, str(_SHARED_DIR))

try:
    from runtime_bootstrap import ensure_plugin_runtime_or_exit
except ImportError:  # pragma: no cover - exercised by isolated script tests
    ensure_plugin_runtime_or_exit = None

if ensure_plugin_runtime_or_exit is not None:
    ensure_plugin_runtime_or_exit(__file__)

from docx import Document
from docx.table import _Cell
from docx.text.paragraph import Paragraph

from pii_common import scan_text_for_pii_matches

@dataclass
class DayPlan:
    date: str
    day_name: str
    standards: list
    learning_intention: str
    success_criteria: list
    agenda: list
    materials: list
    differentiation: list
    evidence: str
    do_now: str = ""


@dataclass
class WeekPlan:
    week_of: str
    subject: str
    teacher: str
    days: list


MAX_PLAN_BYTES = 5 * 1024 * 1024
PLAN_SIDECAR_VERSION = 1

HEADER_MAP = {
    "standards": "standards",
    "gse standards": "standards",
    "learning intention": "learning_intention",
    "learning intentions": "learning_intention",
    "learning target": "learning_intention",
    "success criteria": "success_criteria",
    "i can": "success_criteria",
    "agenda": "agenda",
    "instructional framework": "agenda",
    "lesson": "agenda",
    "materials": "materials",
    "materials needed": "materials",
    "differentiation": "differentiation",
    "differentiated instruction": "differentiation",
    "evidence of learning": "evidence",
    "evaluation of mastery": "evidence",
    "assessment": "evidence",
    "do now": "do_now",
    "do-now": "do_now",
    "bell ringer": "do_now",
    "bellringer": "do_now",
    "warm-up": "do_now",
    "warm up": "do_now",
    "date": "date",
}

PLACEHOLDER_PATTERN = re.compile(r"\{\{[A-Z0-9_]+\}\}")
VALID_DAY_HEADER_RE = re.compile(
    r"^## (?P<date>\d{4}-\d{2}-\d{2})(?:\s*[—-]\s*(?P<day_name>.+))?$"
)
LOOSE_DAY_HEADER_RE = re.compile(r"^## (?P<date>\d{4}-\d{1,2}-\d{1,2})\b")


# ---------------------------------------------------------------------------
# Markdown plan parsing
# ---------------------------------------------------------------------------

def parse_plan(md: str) -> WeekPlan:
    """Parse markdown lesson plan into WeekPlan."""
    lines = md.split("\n")
    week_of = subject = teacher = None
    days: list[DayPlan] = []
    current_day: dict | None = None
    current_section: str | None = None
    section_content: list[str] = []

    def flush(day_dict, section, content):
        if day_dict is None:
            return
        if section:
            day_dict[section] = content
        days.append(_day_from_dict(day_dict))

    for line in lines:
        line_stripped = line.strip()

        if line_stripped.startswith("# Week of"):
            week_of = line_stripped.replace("# Week of", "").strip()
        elif line_stripped.startswith("Subject:"):
            subject = line_stripped.replace("Subject:", "").strip()
        elif line_stripped.startswith("Teacher:"):
            teacher = line_stripped.replace("Teacher:", "").strip()
        elif line_stripped.startswith("## "):
            day_header = VALID_DAY_HEADER_RE.match(line_stripped)
            if day_header:
                flush(current_day, current_section, section_content)
                current_day = {
                    "date": day_header.group("date"),
                    "day_name": (day_header.group("day_name") or "Unknown").strip(),
                }
                current_section = None
                section_content = []
                continue
            loose_header = LOOSE_DAY_HEADER_RE.match(line_stripped)
            if loose_header:
                bad_date = loose_header.group("date")
                raise ValueError(
                    f"Invalid day header '{bad_date}'. "
                    "Use zero-padded YYYY-MM-DD in every '##' day heading."
                )
            continue
        elif line_stripped.startswith("###") and current_day is not None:
            if current_section:
                current_day[current_section] = section_content
            current_section = _normalize_header(
                line_stripped.replace("###", "").strip()
            )
            section_content = []
        elif current_day is not None and current_section is not None:
            if not line_stripped:
                continue
            if line_stripped.startswith("-") or line_stripped.startswith("*"):
                section_content.append(line_stripped[1:].strip())
            elif re.match(r"^\d+\.\s+", line_stripped):
                section_content.append(re.sub(r"^\d+\.\s+", "", line_stripped))
            else:
                section_content.append(line_stripped)

    flush(current_day, current_section, section_content)
    if not days:
        raise ValueError(
            "No lesson days found. Expected day headers like "
            "'## 2026-04-22 — Wednesday'."
        )
    return WeekPlan(week_of or "", subject or "", teacher or "", days)


def _normalize_header(h: str) -> str:
    h_lower = h.lower().strip(":.")
    for key, value in HEADER_MAP.items():
        if key in h_lower:
            return value
    return h_lower.replace(" ", "_")


def _day_from_dict(day_dict: dict) -> DayPlan:
    def _as_list(v):
        if v is None:
            return []
        if isinstance(v, list):
            return v
        return [v]

    def _as_str(v):
        if v is None:
            return ""
        if isinstance(v, list):
            return "\n".join(v)
        return str(v)

    return DayPlan(
        date=day_dict.get("date", ""),
        day_name=day_dict.get("day_name", ""),
        standards=_as_list(day_dict.get("standards")),
        learning_intention=_as_str(day_dict.get("learning_intention")),
        success_criteria=_as_list(day_dict.get("success_criteria")),
        agenda=_as_list(day_dict.get("agenda")),
        materials=_as_list(day_dict.get("materials")),
        differentiation=_as_list(day_dict.get("differentiation")),
        evidence=_as_str(day_dict.get("evidence")),
        do_now=_as_str(day_dict.get("do_now")),
    )


def _compact_text_list(values: list[str]) -> list[str]:
    compacted: list[str] = []
    for value in values:
        text = str(value).strip()
        if text:
            compacted.append(text)
    return compacted


def _build_plan_sidecar_payload(week: WeekPlan) -> dict:
    payload: dict[str, object] = {"v": PLAN_SIDECAR_VERSION, "d": []}
    if week.week_of:
        payload["w"] = week.week_of
    if week.subject:
        payload["s"] = week.subject
    if week.teacher:
        payload["t"] = week.teacher

    day_payloads: list[dict[str, object]] = []
    for day in week.days:
        day_payload: dict[str, object] = {"dt": day.date}
        if day.day_name:
            day_payload["n"] = day.day_name
        if day.standards:
            day_payload["st"] = _compact_text_list(day.standards)
        if day.learning_intention:
            day_payload["li"] = day.learning_intention
        if day.success_criteria:
            day_payload["sc"] = _compact_text_list(day.success_criteria)
        if day.agenda:
            day_payload["ag"] = _compact_text_list(day.agenda)
        if day.materials:
            day_payload["m"] = _compact_text_list(day.materials)
        if day.differentiation:
            day_payload["df"] = _compact_text_list(day.differentiation)
        if day.evidence:
            day_payload["e"] = day.evidence
        if day.do_now:
            day_payload["do"] = day.do_now
        day_payloads.append(day_payload)

    payload["d"] = day_payloads
    return payload


def _serialize_sidecar_json(payload: dict) -> str:
    return json.dumps(payload, separators=(",", ":"), sort_keys=True)


def _existing_json_sidecar_matches(sidecar_path: Path, expected_payload: dict) -> bool:
    try:
        existing_text = sidecar_path.read_text(encoding="utf-8")
    except OSError as exc:
        raise OSError(
            f"could not read existing JSON sidecar {sidecar_path}: {exc}"
        ) from exc
    try:
        existing_payload = json.loads(existing_text)
    except json.JSONDecodeError as exc:
        raise ValueError(
            f"existing JSON sidecar {sidecar_path} is invalid: {exc}"
        ) from exc
    return existing_payload == expected_payload


# ---------------------------------------------------------------------------
# PII scanning — layered: keyword patterns + bare-name detector
# ---------------------------------------------------------------------------

def scan_for_pii(text: str, allowed_names: set[str]) -> list[tuple[str, str]]:
    """Scan text for PII. Returns ``[(matched_text, pattern_label), ...]``."""
    return scan_text_for_pii_matches(text, allowed_names)


def scan_docx_content(doc: Document, allowed_names: set[str]) -> list[tuple[str, str]]:
    """Scan a loaded docx's visible text for PII.

    Covers the BODY (paragraphs, tables, and text boxes reachable via
    ``w:p`` descent — ``iter(w:p)`` walks into ``w:txbxContent``) AND
    every section's headers and footers. Pre-fix this only scanned
    ``doc.paragraphs`` + ``doc.tables``, so a district template with a
    header that read ``Student: John Doe`` wrote clean to disk.
    """
    try:
        pieces = [
            "".join(run.text for run in para.runs)
            for para in _iter_docx_paragraphs(doc)
            if para.runs
        ]
    except Exception:
        # python-docx is a hard dep for this script — if its internals
        # aren't available, fail closed by returning a fake match so
        # the caller aborts rather than silently skipping the scan.
        return [("<docx scan unavailable>", "scan error")]

    return scan_for_pii("\n".join(pieces), allowed_names)


def _iter_docx_story_paragraphs(doc: Document) -> Iterable[Paragraph]:
    """Yield paragraphs across body, tables, headers/footers, and text boxes."""
    from docx.oxml.ns import qn

    yield_parent_pairs = []
    body = getattr(doc.element, "body", None)
    if body is not None:
        yield_parent_pairs.append((body, doc))

    seen_ids: set[int] = set()
    for section in doc.sections:
        for attr in (
            "header", "first_page_header", "even_page_header",
            "footer", "first_page_footer", "even_page_footer",
        ):
            hdr_ftr = getattr(section, attr, None)
            if hdr_ftr is None:
                continue
            element = getattr(hdr_ftr, "_element", None)
            if element is None:
                part = getattr(hdr_ftr, "part", None)
                element = getattr(part, "element", None) if part is not None else None
            if element is None or id(element) in seen_ids:
                continue
            seen_ids.add(id(element))
            yield_parent_pairs.append((element, hdr_ftr))

    for root, parent in yield_parent_pairs:
        for p in list(root.iter(qn("w:p"))):
            yield Paragraph(p, parent)


def _iter_docx_paragraphs(doc: Document) -> Iterable[Paragraph]:
    yield from _iter_docx_story_paragraphs(doc)


def _find_unresolved_docx_placeholders(doc: Document) -> list[str]:
    found: set[str] = set()
    for para in _iter_docx_paragraphs(doc):
        text = "".join(run.text for run in para.runs if run.text)
        found.update(PLACEHOLDER_PATTERN.findall(text))
    return sorted(found)


# ---------------------------------------------------------------------------
# docx formatting-preserving writes
# ---------------------------------------------------------------------------

def _set_cell_text_preserving_format(cell: _Cell, value: str) -> None:
    """Write value into the cell's first paragraph's first run, preserving
    run-level formatting (font, bold, color, size). Clears any additional
    runs and extra paragraphs. Multi-line values become multi-paragraph
    writes cloned from the first paragraph's style.
    """
    if not cell.paragraphs:
        cell.add_paragraph(str(value))
        return

    lines = str(value).splitlines() or [""]

    # Clear paragraphs beyond what we need by blanking their text (pure
    # python-docx doesn't offer element removal without private APIs; we
    # prefer a non-destructive approach here).
    for para in cell.paragraphs[1:]:
        for run in para.runs:
            run.text = ""

    first_para = cell.paragraphs[0]
    _set_paragraph_text_preserving_format(first_para, lines[0])

    # Additional lines — add new paragraphs copying the first paragraph's
    # style. python-docx doesn't clone run formatting across add_paragraph,
    # so we replicate by copying the first run's properties onto a new run.
    for line in lines[1:]:
        new_para = cell.add_paragraph()
        if first_para.style is not None:
            try:
                new_para.style = first_para.style
            except Exception:
                pass
        run_source = first_para.runs[0] if first_para.runs else None
        new_run = new_para.add_run(line)
        if run_source is not None:
            _copy_run_format(run_source, new_run)


def _set_paragraph_text_preserving_format(para, value: str) -> None:
    """Replace a paragraph's visible text without destroying run formatting."""
    if not para.runs:
        para.add_run(value)
        return
    # Put all text in the first run, blank all trailing runs. This keeps
    # font/bold/italic/color/hyperlink definitions intact on the first run.
    para.runs[0].text = value
    for run in para.runs[1:]:
        run.text = ""


def _copy_run_format(src, dst) -> None:
    """Copy common run formatting from src to dst."""
    try:
        if src.bold is not None:
            dst.bold = src.bold
        if src.italic is not None:
            dst.italic = src.italic
        if src.underline is not None:
            dst.underline = src.underline
        if src.font.name:
            dst.font.name = src.font.name
        if src.font.size:
            dst.font.size = src.font.size
        if src.font.color and src.font.color.rgb:
            dst.font.color.rgb = src.font.color.rgb
    except Exception:
        # Defensive — python-docx raises on some unreachable color states.
        pass


def _replace_in_paragraph(para, replacements: dict[str, str]) -> bool:
    """Replace {{PLACEHOLDER}} tokens in a paragraph across runs, preserving
    formatting on the run that held the opening `{{`.

    python-docx splits placeholders across runs surprisingly often — we
    concatenate run text, do the substitution, then push the result back
    into the first run (clearing the rest). Formatting on the first run
    survives; formatting on runs we blank would have been lost anyway.
    """
    text_runs = [run for run in para.runs if run.text]
    if not text_runs:
        return False
    combined = "".join(run.text for run in text_runs)
    changed = False
    for key, value in replacements.items():
        placeholder = f"{{{{{key}}}}}"
        if placeholder in combined:
            combined = combined.replace(placeholder, str(value))
            changed = True
    if changed:
        text_runs[0].text = combined
        for run in text_runs[1:]:
            run.text = ""
    return changed


# ---------------------------------------------------------------------------
# Fill strategies
# ---------------------------------------------------------------------------

def fill_placeholders(doc: Document, week: WeekPlan) -> bool:
    """Fill explicit {{PLACEHOLDERS}} throughout the document."""
    replacements: dict[str, str] = {
        "WEEK_OF": week.week_of,
        "SUBJECT": week.subject,
        "TEACHER": week.teacher,
    }

    for day in week.days:
        day_upper = (day.day_name or "").upper()
        if not day_upper:
            continue
        replacements[f"{day_upper}_DATE"] = day.date
        replacements[f"{day_upper}_STANDARDS"] = "\n".join(day.standards)
        replacements[f"{day_upper}_LEARNING_INTENTION"] = day.learning_intention
        replacements[f"{day_upper}_SUCCESS_CRITERIA"] = "\n".join(day.success_criteria)
        replacements[f"{day_upper}_AGENDA"] = "\n".join(day.agenda)
        replacements[f"{day_upper}_MATERIALS"] = "\n".join(day.materials)
        replacements[f"{day_upper}_DIFFERENTIATION"] = "\n".join(day.differentiation)
        replacements[f"{day_upper}_EVIDENCE"] = day.evidence
        replacements[f"{day_upper}_DO_NOW"] = day.do_now

    found_any = False
    for para in _iter_docx_paragraphs(doc):
        if _replace_in_paragraph(para, replacements):
            found_any = True
    return found_any


_DAY_NAME_PATTERNS = {
    "monday": "Monday", "mon": "Monday",
    "tuesday": "Tuesday", "tues": "Tuesday", "tue": "Tuesday",
    "wednesday": "Wednesday", "wed": "Wednesday",
    "thursday": "Thursday", "thurs": "Thursday", "thu": "Thursday",
    "friday": "Friday", "fri": "Friday",
}


def _maybe_day_from_cell(text: str) -> str | None:
    low = text.lower().strip(":. \t\n")
    if low in _DAY_NAME_PATTERNS:
        return _DAY_NAME_PATTERNS[low]
    for key, canonical in _DAY_NAME_PATTERNS.items():
        if low.startswith(key):
            return canonical
    return None


def _header_field_for(text: str) -> str | None:
    low = text.lower().strip(":.")
    for header_key, field_name in HEADER_MAP.items():
        if header_key in low:
            return field_name
    return None


def fill_by_headers(
    doc: Document, week: WeekPlan
) -> tuple[bool, list[str]]:
    """Fill tables by header-matching.

    Handles both template shapes:
      1. Per-table cardinality: one table per day, with rows of
         (field_label, value) pairs. Day index advances per table.
      2. Per-row cardinality: one table with a column-per-day or a
         row-per-day, where the leftmost cell names the day.

    Returns (filled_something, unmatched_notes).
    """
    unmatched: list[str] = []
    filled_any = False
    days_by_name: dict[str, DayPlan] = {
        (d.day_name or "").lower(): d for d in week.days
    }

    for t_idx, table in enumerate(doc.tables):
        if not table.rows:
            continue

        # Detect per-row cardinality: look at the first column for day names.
        first_col_days: list[tuple[int, str]] = []
        for r_idx, row in enumerate(table.rows):
            if not row.cells:
                continue
            first_cell = row.cells[0]
            day_candidate = _maybe_day_from_cell(first_cell.text)
            if day_candidate:
                first_col_days.append((r_idx, day_candidate))

        # Detect per-column cardinality: row 0 holds day names across columns.
        header_row = table.rows[0]
        col_day_map: dict[int, str] = {}
        if header_row.cells:
            for c_idx, cell in enumerate(header_row.cells):
                day_candidate = _maybe_day_from_cell(cell.text)
                if day_candidate:
                    col_day_map[c_idx] = day_candidate

        # --- Shape A: day-per-row ---
        if len(first_col_days) >= 2:
            for r_idx, day_name in first_col_days:
                day = days_by_name.get(day_name.lower())
                if not day:
                    unmatched.append(
                        f"Table {t_idx}: row {r_idx} labels day "
                        f"'{day_name}' but plan has no matching day."
                    )
                    continue
                row = table.rows[r_idx]
                # Each remaining cell in the row is a field. We map
                # them by looking at the header row's column labels.
                for c_idx, cell in enumerate(row.cells[1:], start=1):
                    field_name = None
                    if header_row and c_idx < len(header_row.cells):
                        field_name = _header_field_for(header_row.cells[c_idx].text)
                    if not field_name:
                        # Try the cell itself for a "Field: value" pattern
                        continue
                    value = getattr(day, field_name, "")
                    if isinstance(value, list):
                        value = "\n".join(value)
                    _set_cell_text_preserving_format(cell, str(value))
                    filled_any = True
            continue

        # --- Shape B: day-per-column (header row names days) ---
        if len(col_day_map) >= 2:
            # Track missing days once per table rather than once per field row.
            missing_days: set[str] = set()
            for r_idx, row in enumerate(table.rows):
                if r_idx == 0:
                    continue
                if not row.cells:
                    continue
                field_name = _header_field_for(row.cells[0].text)
                if not field_name:
                    continue
                for c_idx, day_name in col_day_map.items():
                    if c_idx >= len(row.cells):
                        continue
                    day = days_by_name.get(day_name.lower())
                    if not day:
                        missing_days.add(day_name)
                        continue
                    value = getattr(day, field_name, "")
                    if isinstance(value, list):
                        value = "\n".join(value)
                    _set_cell_text_preserving_format(row.cells[c_idx], str(value))
                    filled_any = True
            for day_name in sorted(missing_days):
                unmatched.append(
                    f"Table {t_idx}: header column '{day_name}' "
                    f"has no matching day in plan (left blank)."
                )
            continue

        # --- Shape C: one table per day, field-label rows ---
        if t_idx < len(week.days):
            day = week.days[t_idx]
            table_filled = False
            for r_idx, row in enumerate(table.rows):
                if len(row.cells) < 2:
                    continue
                field_name = _header_field_for(row.cells[0].text)
                if not field_name:
                    continue
                value = getattr(day, field_name, "")
                if isinstance(value, list):
                    value = "\n".join(value)
                _set_cell_text_preserving_format(row.cells[1], str(value))
                filled_any = True
                table_filled = True
            if not table_filled:
                unmatched.append(
                    f"Table {t_idx}: no field labels matched HEADER_MAP "
                    f"for day '{day.day_name}'."
                )

    return filled_any, unmatched


# ---------------------------------------------------------------------------
# Config I/O — privacy validation + opt-in mutations
# ---------------------------------------------------------------------------

def load_config(config_path: Path) -> dict:
    """Load the teacher config YAML. Returns empty dict if missing."""
    if not config_path.exists():
        return {}
    try:
        import yaml
    except ImportError:
        return {}
    try:
        with open(config_path, "r", encoding="utf-8") as f:
            data = yaml.safe_load(f)
    except OSError as exc:
        raise ValueError(f"Could not read config file: {exc}") from exc
    except Exception as exc:
        raise ValueError(f"Could not load config: {exc}") from exc
    if data is None:
        return {}
    if not isinstance(data, dict):
        raise ValueError("Config file must decode to a YAML mapping.")
    return data


def validate_privacy(cfg: dict) -> list[str]:
    """Return a list of privacy-invariant violations (empty = all pass).

    PyYAML uses the YAML 1.1 core schema where ``off``, ``no``, ``false``
    all parse to ``False`` and ``on``, ``yes``, ``true`` to ``True``.
    A teacher who writes the documented ``telemetry: off`` without
    quotes gets ``False`` — which is the correct intent, not a
    misconfiguration. We accept that, while still rejecting ambiguous
    values like ``"no"`` (string), ``0``, or ``null`` for either field.
    ``student_data`` has no correct boolean form (``"never"`` isn't a
    YAML boolean) — only the literal string is accepted.
    """
    errors: list[str] = []
    priv = (cfg or {}).get("privacy", {}) or {}

    student_data = priv.get("student_data", "never")
    # "never" is not a YAML boolean, so we don't accept any boolean
    # substitute here. Only the literal string "never" (any case).
    if not isinstance(student_data, str) or student_data.lower() != "never":
        actual = repr(student_data)
        errors.append(
            f"privacy.student_data must be the string 'never' "
            f"(got {actual}; booleans and other types are not accepted)"
        )

    telemetry = priv.get("telemetry", "off")
    # Accept the YAML 1.1 boolean False (parsed from ``off``, ``no``,
    # or ``false``) AND the literal string "off" (any case). Reject
    # everything else — e.g. integers, None, True, or strings like "no".
    telemetry_ok = (
        telemetry is False
        or (isinstance(telemetry, str) and telemetry.lower() == "off")
    )
    if not telemetry_ok:
        actual = repr(telemetry)
        errors.append(
            f"privacy.telemetry must be 'off' (or the YAML keyword "
            f"``off`` which parses as the boolean false); got {actual}"
        )
    # Recommended but not hard — warn-only by returning None-type tags
    # would be fine; we keep scan_before_write a hard invariant because
    # the audit's D6 asked for enforcement.
    scan_before_write = priv.get("pii_scan_before_write", True)
    if scan_before_write is False:
        errors.append(
            "privacy.pii_scan_before_write: true is required "
            "(turning off the scan defeats the plugin's FERPA guarantee)"
        )
    elif scan_before_write is not True:
        errors.append(
            "privacy.pii_scan_before_write must be the boolean true when "
            "present; quoted strings like \"false\" or \"true\" are not accepted"
        )
    return errors


def write_config_mutation(
    config_path: Path,
    subject_id: str,
    mutation: dict,
) -> bool:
    """Update subjects[id==subject_id] with the given mutation dict.

    Prefers ruamel.yaml to preserve comments; falls back to PyYAML.
    Returns True on successful write.
    """
    try:
        config_path = _normalize_config_path(config_path)
    except ValueError as exc:
        print(f"Error: {exc}", file=sys.stderr)
        return False

    if not config_path.exists():
        print(
            f"Warning: config not found at {config_path}; skipping mutation.",
            file=sys.stderr,
        )
        return False

    data = None
    ruamel_available = False
    try:
        from ruamel.yaml import YAML
        yaml_rt = YAML()
        yaml_rt.preserve_quotes = True
        with open(config_path, "r", encoding="utf-8") as f:
            data = yaml_rt.load(f)
        ruamel_available = True
    except Exception:
        try:
            import yaml
            with open(config_path, "r", encoding="utf-8") as f:
                data = yaml.safe_load(f) or {}
        except Exception as e:
            print(f"Error: failed to load config: {e}", file=sys.stderr)
            return False

    if not data or "subjects" not in data:
        print(
            "Warning: config has no 'subjects' block; skipping mutation.",
            file=sys.stderr,
        )
        return False

    target = None
    for subject in data["subjects"]:
        if subject.get("id") == subject_id:
            target = subject
            break
    if target is None:
        print(
            f"Warning: subject id '{subject_id}' not found in config; "
            f"skipping mutation.",
            file=sys.stderr,
        )
        return False

    # Apply mutation (dotted keys like template.mapping_verified -> nested)
    for dotted_key, value in mutation.items():
        parts = dotted_key.split(".")
        node = target
        for part in parts[:-1]:
            if part not in node or not isinstance(node[part], dict):
                node[part] = {}
            node = node[part]
        node[parts[-1]] = value

    try:
        if ruamel_available:
            with open(config_path, "w", encoding="utf-8") as f:
                yaml_rt.dump(data, f)
        else:
            import yaml
            with open(config_path, "w", encoding="utf-8") as f:
                yaml.safe_dump(data, f, sort_keys=False)
        return True
    except Exception as e:
        print(f"Error: failed to write config: {e}", file=sys.stderr)
        return False


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def _home_root() -> Path:
    # Override with LESSON_PLAN_MAGIC_HOME for managed installs (e.g. Windows
    # profiles where Documents is redirected) or non-standard layouts.
    override = os.environ.get("LESSON_PLAN_MAGIC_HOME")
    if override:
        return Path(override).expanduser()
    return Path.home() / "Documents" / "Lesson Plan Magic"


def _default_config_path() -> Path:
    return _home_root() / "config.yaml"


def _config_root() -> Path:
    return _home_root().resolve()


def _output_root() -> Path:
    return (_config_root() / "outputs").resolve()


def _normalize_config_path(config_path: Path) -> Path:
    resolved = config_path.expanduser().resolve()
    root = _config_root()
    try:
        resolved.relative_to(root)
    except ValueError as exc:
        raise ValueError(
            f"Config path must stay within {root}; got {resolved}"
        ) from exc
    return resolved


def _normalize_output_path(
    output_path: Path, *, allow_anywhere: bool = False
) -> Path:
    resolved = output_path.expanduser().resolve()
    if allow_anywhere:
        return resolved

    root = _output_root()
    try:
        resolved.relative_to(root)
    except ValueError as exc:
        raise ValueError(
            f"Output path must stay within {root} unless --allow-anywhere is passed."
        ) from exc
    return resolved


def _read_plan_markdown(plan_path: Path) -> str:
    try:
        size = plan_path.stat().st_size
    except OSError as exc:
        raise ValueError(f"Could not stat plan file: {exc}") from exc
    if size > MAX_PLAN_BYTES:
        raise ValueError(
            f"Plan markdown exceeds the {MAX_PLAN_BYTES}-byte limit: {plan_path}"
        )
    return plan_path.read_text(encoding="utf-8")


def main() -> int:
    parser = argparse.ArgumentParser(
        description="Fill lesson-plan template from markdown."
    )
    parser.add_argument("--plan", required=True, help="Path to markdown lesson plan")
    parser.add_argument("--template", required=True, help="Path to .docx template")
    parser.add_argument("--output", required=True, help="Output .docx path")
    parser.add_argument(
        "--allow-names",
        help=(
            "Comma-separated bare-name allowlist "
            "(teacher/public figures only; does not bypass Student:/IEP:/email/phone hits)"
        ),
    )
    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="Parse and scan only; do not write.",
    )
    parser.add_argument(
        "--config",
        help="Path to teacher config.yaml "
        "(default: ~/Documents/Lesson Plan Magic/config.yaml).",
    )
    parser.add_argument(
        "--allow-anywhere",
        action="store_true",
        help="Allow writing outside ~/Documents/Lesson Plan Magic/outputs/.",
    )
    parser.add_argument(
        "--mark-verified",
        metavar="SUBJECT_ID",
        help="After a successful fill, set subjects[<id>].template.mapping_verified=true.",
    )
    parser.add_argument(
        "--update-pacing",
        metavar="SUBJECT_ID",
        help="After a successful fill, set subjects[<id>].pacing.last_planned_week_end "
        "to the last day's date in the plan.",
    )
    parser.add_argument(
        "--skip-docx-scan",
        action="store_true",
        help=(
            "Skip scanning the loaded template's own text for PII "
            "(debugging escape hatch for a known-clean legacy template)"
        ),
    )
    parser.add_argument(
        "--no-sidecar",
        action="store_true",
        help="Skip writing plan sidecars .plan.md and .plan.json (default: sidecars are written).",
    )

    args = parser.parse_args()

    plan_path = Path(args.plan)
    template_path = Path(args.template)
    output_path = Path(args.output)
    try:
        config_path = _normalize_config_path(
            Path(args.config).expanduser() if args.config else _default_config_path()
        )
    except ValueError as exc:
        print(f"Error: {exc}", file=sys.stderr)
        return 1
    try:
        output_path = _normalize_output_path(
            output_path, allow_anywhere=args.allow_anywhere
        )
    except ValueError as exc:
        print(f"Error: {exc}", file=sys.stderr)
        return 1

    if not plan_path.exists():
        print(f"Error: Plan file not found: {plan_path}", file=sys.stderr)
        return 1
    if not template_path.exists():
        print(f"Error: Template file not found: {template_path}", file=sys.stderr)
        return 1

    # Validate privacy invariants if a config is present.
    try:
        cfg = load_config(config_path)
    except ValueError as exc:
        print(f"Error: {exc}", file=sys.stderr)
        return 1
    if cfg:
        errors = validate_privacy(cfg)
        if errors:
            for e in errors:
                print(f"Config privacy error: {e}", file=sys.stderr)
            return 3

    # Merge allowed-names from CLI + config.
    #
    # A teacher's own name is not student PII. The FERPA-safe guarantee
    # covers student identifying data — the teacher's name is on the
    # plan header, the template ({{TEACHER}}), the sub-plan, the exit
    # ticket, the email sig. We auto-allowlist it from both the config
    # (teacher.name) and the parsed plan (Teacher: line) so routine
    # teacher data doesn't trip the scanner. See hard rule #1 in SKILL.md
    # — it targets *student* names, not the teacher's.
    allowed_names: set[str] = set()
    if args.allow_names:
        allowed_names.update(n.strip() for n in args.allow_names.split(",") if n.strip())
    if cfg.get("approved_names"):
        allowed_names.update(
            n.strip()
            for n in cfg["approved_names"]
            if isinstance(n, str) and n.strip()
        )
    teacher_cfg_name = (cfg.get("teacher") or {}).get("name")
    if isinstance(teacher_cfg_name, str) and teacher_cfg_name.strip():
        allowed_names.add(teacher_cfg_name.strip())
    # Also exempt a co-teacher name if the active subject has one.
    for subject in cfg.get("subjects") or []:
        co = subject.get("co_teacher_name") if isinstance(subject, dict) else None
        if isinstance(co, str) and co.strip():
            allowed_names.add(co.strip())

    try:
        md_content = _read_plan_markdown(plan_path)
    except ValueError as exc:
        print(f"Error: {exc}", file=sys.stderr)
        return 1
    try:
        week = parse_plan(md_content)
    except ValueError as exc:
        print(f"Error: {exc}", file=sys.stderr)
        return 1

    # The plan's own "Teacher:" line is not PII either — allowlist it
    # before scanning. parse_plan already stripped the label, so
    # week.teacher is just the name.
    if week.teacher and week.teacher.strip():
        allowed_names.add(week.teacher.strip())

    # Scan the raw markdown, not only the parsed canonical fields. The sidecar
    # writes the original markdown verbatim, so unrecognized sections like
    # "### Notes" must be privacy-gated too.
    pii_matches = scan_for_pii(md_content, allowed_names)
    if pii_matches:
        for _matched_text, pattern_name in pii_matches:
            print(
                f"PII detected in plan: {pattern_name}",
                file=sys.stderr,
            )
        return 2

    if args.dry_run:
        print("Dry-run: PII scan passed. Would write output.", file=sys.stderr)
        return 0

    doc = Document(str(template_path))

    # Scan the template itself (catches district rosters baked into the
    # template, which the plan-text scan would miss).
    if not args.skip_docx_scan:
        tpl_hits = scan_docx_content(doc, allowed_names)
        if tpl_hits:
            for _matched_text, label in tpl_hits:
                print(
                    f"PII detected in template: {label}",
                    file=sys.stderr,
                )
            print(
                "Aborting: clean the template or allowlist historical names "
                "via --allow-names.",
                file=sys.stderr,
            )
            return 2

    # Strategy 1: explicit placeholders (covers meta fields like
    # {{TEACHER}} / {{SUBJECT}} / {{WEEK_OF}} and any day-specific
    # placeholders the template happens to carry).
    placeholders_hit = fill_placeholders(doc, week)

    # Strategy 2: header-matching. Always run — starter templates have a
    # meta row with placeholders AND a table that needs header-fill, so
    # the two strategies are complementary rather than mutually exclusive.
    # fill_by_headers is a no-op on tables with no matching labels.
    header_hit, unmatched = fill_by_headers(doc, week)

    if not placeholders_hit and not header_hit:
        print(
            "Error: could not map the plan to the template — "
            "no {{PLACEHOLDERS}} found and no header cells matched HEADER_MAP.",
            file=sys.stderr,
        )
        print(
            "Hard rule #3: this script does not silently append fresh tables.",
            file=sys.stderr,
        )
        print(
            "Fix: add {{PLACEHOLDER}} tokens to the template, "
            "or rename header cells (see HEADER_MAP for recognized labels), "
            "or pass --template pointing to a starter template.",
            file=sys.stderr,
        )
        return 4

    if unmatched:
        print(
            "Warning: some days/fields did not match the template:",
            file=sys.stderr,
        )
        for note in unmatched:
            print(f"  - {note}", file=sys.stderr)

    unresolved = _find_unresolved_docx_placeholders(doc)
    if unresolved:
        print(
            "Error: unresolved template placeholders remain after fill:",
            file=sys.stderr,
        )
        for token in unresolved[:10]:
            print(f"  - {token}", file=sys.stderr)
        if len(unresolved) > 10:
            print(f"  - ... and {len(unresolved) - 10} more", file=sys.stderr)
        print(
            "Fix the template mapping or remove unsupported placeholders "
            "before retrying.",
            file=sys.stderr,
        )
        return 4

    # Final output-side scan (belt + suspenders — catches any placeholder
    # that pulled in ambient text we didn't anticipate).
    final_hits = scan_docx_content(doc, allowed_names)
    if final_hits:
        for _matched_text, label in final_hits:
            print(
                f"PII detected in filled output: {label}",
                file=sys.stderr,
            )
        return 2

    md_sidecar_path = None
    md_sidecar_unchanged = False
    json_sidecar_path = None
    json_sidecar_unchanged = False
    plan_sidecar_payload = _build_plan_sidecar_payload(week)
    plan_sidecar_json = _serialize_sidecar_json(plan_sidecar_payload)
    if not args.no_sidecar:
        md_sidecar_path = output_path.with_suffix(".plan.md")
        if md_sidecar_path.exists():
            try:
                existing = md_sidecar_path.read_text(encoding="utf-8")
            except OSError as exc:
                print(
                    f"Error: could not read existing sidecar {md_sidecar_path}: {exc}",
                    file=sys.stderr,
                )
                return 5
            if existing != md_content:
                print(
                    f"Error: refusing to overwrite existing sidecar {md_sidecar_path}. "
                    "Delete it first if replacement is intentional.",
                    file=sys.stderr,
                )
                return 5
            md_sidecar_unchanged = True

        json_sidecar_path = output_path.with_suffix(".plan.json")
        if json_sidecar_path.exists():
            try:
                matches_existing = _existing_json_sidecar_matches(
                    json_sidecar_path, plan_sidecar_payload
                )
            except OSError as exc:
                print(f"Error: {exc}", file=sys.stderr)
                return 5
            except ValueError as exc:
                print(f"Error: {exc}", file=sys.stderr)
                return 5
            if not matches_existing:
                print(
                    f"Error: refusing to overwrite existing JSON sidecar {json_sidecar_path}. "
                    "Delete it first if replacement is intentional.",
                    file=sys.stderr,
                )
                return 5
            json_sidecar_unchanged = True

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Wrote: {output_path}")

    # Write markdown sidecar (private bridge to classroom-artifacts skill).
    # Only written after final PII scan passes — guarantees we never persist
    # a plan that failed privacy. Fail closed if a different sidecar already
    # exists; callers must intentionally remove or rename it first.
    if not args.no_sidecar:
        if md_sidecar_unchanged:
            print(f"Sidecar unchanged: {md_sidecar_path}", file=sys.stderr)
        else:
            md_sidecar_path.write_text(md_content, encoding="utf-8")
            print(f"Wrote sidecar: {md_sidecar_path}", file=sys.stderr)
        if json_sidecar_unchanged:
            print(
                f"JSON sidecar unchanged: {json_sidecar_path}",
                file=sys.stderr,
            )
        else:
            json_sidecar_path.write_text(plan_sidecar_json, encoding="utf-8")
            print(f"Wrote JSON sidecar: {json_sidecar_path}", file=sys.stderr)

    # Optional config mutations
    if args.mark_verified:
        ok = write_config_mutation(
            config_path,
            args.mark_verified,
            {"template.mapping_verified": True},
        )
        if ok:
            print(
                f"Updated config: subjects[{args.mark_verified}]"
                f".template.mapping_verified = true",
                file=sys.stderr,
            )

    if args.update_pacing and week.days:
        last_date = week.days[-1].date
        if re.match(r"^\d{4}-\d{2}-\d{2}$", last_date or ""):
            ok = write_config_mutation(
                config_path,
                args.update_pacing,
                {"pacing.last_planned_week_end": last_date},
            )
            if ok:
                print(
                    f"Updated config: subjects[{args.update_pacing}]"
                    f".pacing.last_planned_week_end = {last_date}",
                    file=sys.stderr,
                )

    return 0


if __name__ == "__main__":
    sys.exit(main())
