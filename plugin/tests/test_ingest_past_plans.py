"""Tests for ingest_past_plans.py — .docx extraction, PII redaction,
signature extractors, pacing, markdown formatter, and CLI."""
from __future__ import annotations

import json
import subprocess
import sys
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

import pytest
from docx import Document

import ingest_past_plans as ipp


SCRIPT_PATH = (
    Path(__file__).parent.parent
    / "skills" / "lesson-planner" / "scripts" / "ingest_past_plans.py"
)
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


# ---------------------------------------------------------------------------
# Fixture builders
# ---------------------------------------------------------------------------

def _build_docx(path: Path, paragraphs=(), tables=()) -> None:
    """Write a .docx with the given paragraphs and optional tables.

    `tables` is a list of row-lists, each row a list of cell strings.
    """
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    for rows in tables:
        if not rows:
            continue
        tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
        for r, row in enumerate(rows):
            for c, val in enumerate(row):
                tbl.rows[r].cells[c].text = val
    doc.save(str(path))


def _rewrite_docx_part(path: Path, member: str, transform) -> None:
    tmp_path = path.with_suffix(path.suffix + ".tmp")
    found = False
    with zipfile.ZipFile(path) as src, zipfile.ZipFile(tmp_path, "w") as dst:
        for info in src.infolist():
            data = src.read(info.filename)
            if info.filename == member:
                data = transform(data)
                found = True
            dst.writestr(info, data)
        if not found:
            dst.writestr(member, transform(None))
    tmp_path.replace(path)


def _inject_tracked_change_author(path: Path, author: str) -> None:
    def transform(data: bytes | None) -> bytes:
        if data is None:
            raise AssertionError("word/document.xml missing from synthetic test docx")
        root = ET.fromstring(data)
        body = root.find(f"{{{W_NS}}}body")
        assert body is not None

        paragraph = ET.Element(f"{{{W_NS}}}p")
        change = ET.SubElement(
            paragraph,
            f"{{{W_NS}}}ins",
            {
                f"{{{W_NS}}}id": "1",
                f"{{{W_NS}}}author": author,
                f"{{{W_NS}}}date": "2026-04-21T00:00:00Z",
            },
        )
        run = ET.SubElement(change, f"{{{W_NS}}}r")
        text = ET.SubElement(run, f"{{{W_NS}}}t")
        text.text = "Tracked revision"
        body.insert(0, paragraph)
        return ET.tostring(root, encoding="utf-8", xml_declaration=True)

    _rewrite_docx_part(path, "word/document.xml", transform)


# ---------------------------------------------------------------------------
# extract_text_from_docx
# ---------------------------------------------------------------------------

class TestExtractTextFromDocx:
    def test_paragraphs_and_tables(self, tmp_path):
        f = tmp_path / "plan.docx"
        _build_docx(
            f,
            paragraphs=["Lesson: photosynthesis", "Warm-up: think-pair-share"],
            tables=[[["Objective", "Task"], ["Explain ATP", "Gallery walk"]]],
        )
        text, tables, rows = ipp.extract_text_from_docx(f)
        assert "photosynthesis" in text
        assert "Gallery walk" in text
        assert tables == 1
        assert rows == 2

    def test_empty_doc(self, tmp_path):
        f = tmp_path / "empty.docx"
        _build_docx(f)
        text, tables, rows = ipp.extract_text_from_docx(f)
        assert text == ""
        assert tables == 0
        assert rows == 0

    def test_skips_blank_paragraphs_and_cells(self, tmp_path):
        f = tmp_path / "plan.docx"
        _build_docx(
            f,
            paragraphs=["real content", "   ", ""],
            tables=[[["", "keep"], ["  ", "also"]]],
        )
        text, tables, rows = ipp.extract_text_from_docx(f)
        assert text.count("\n") < 5  # blanks got stripped
        assert "real content" in text and "keep" in text and "also" in text
        assert tables == 1 and rows == 2

    def test_unreadable_file_returns_empty(self, tmp_path):
        f = tmp_path / "broken.docx"
        f.write_text("this is not a real .docx", encoding="utf-8")
        text, tables, rows = ipp.extract_text_from_docx(f)
        assert text == "" and tables == 0 and rows == 0

    def test_oversized_docx_xml_part_returns_empty(self, tmp_path):
        f = tmp_path / "bomb.docx"
        with zipfile.ZipFile(f, "w") as zf:
            zf.writestr(
                "word/document.xml",
                b"x" * (ipp.MAX_DOCX_XML_ENTRY_BYTES + 1),
            )

        text, tables, rows = ipp.extract_text_from_docx(f)
        assert text == ""
        assert tables == 0
        assert rows == 0

    def test_metadata_author_extraction_ignores_oversized_xml_part(self, tmp_path):
        f = tmp_path / "bomb.docx"
        with zipfile.ZipFile(f, "w") as zf:
            zf.writestr(
                "word/comments.xml",
                b"x" * (ipp.MAX_DOCX_XML_ENTRY_BYTES + 1),
            )

        assert ipp.extract_docx_metadata_authors(f) == []


# ---------------------------------------------------------------------------
# redact_names — all 6 redaction passes
# ---------------------------------------------------------------------------

class TestRedactNames:
    def test_redact_disabled(self):
        text = "John Smith attended class."
        out, n = ipp.redact_names(text, redact=False)
        assert out == text and n == 0

    # Pass 1: Firstname Lastname
    def test_full_name_redacted(self):
        out, n = ipp.redact_names("John Smith was absent.")
        assert "John Smith" not in out
        assert "[STUDENT]" in out
        assert n == 1

    def test_historical_allowlist_preserved(self):
        text = "We discussed Rosa Parks and Abraham Lincoln."
        out, n = ipp.redact_names(text)
        assert "Rosa Parks" in out and "Abraham Lincoln" in out
        assert n == 0

    def test_teacher_name_allowlisted(self):
        out, n = ipp.redact_names(
            "Ms. Jane Doe will lead the activity.",
            teacher_name="Jane Doe",
        )
        assert "Jane Doe" in out
        assert n == 0

    # Pass 2: initials
    def test_initials_two(self):
        out, n = ipp.redact_names("Student A.B. needs extended time.")
        assert "A.B." not in out
        assert "[STUDENT]" in out
        assert n >= 1

    def test_initials_three(self):
        out, n = ipp.redact_names("Tracking progress for A.B.C. this week.")
        assert "A.B.C." not in out
        assert n >= 1

    def test_initials_with_space(self):
        out, n = ipp.redact_names("Reminder about J. S. during 3rd period.")
        assert "J. S." not in out
        assert n >= 1

    # Pass 3: emails
    def test_email_redacted(self):
        out, n = ipp.redact_names("Contact jsmith@school.edu for details.")
        assert "jsmith@school.edu" not in out
        assert "[EMAIL]" in out
        assert n >= 1

    # Pass 4: student IDs
    def test_student_id_with_prefix(self):
        out, n = ipp.redact_names("Student #1234567 should see nurse.")
        assert "1234567" not in out
        assert "[STUDENT-ID]" in out
        assert n >= 1

    def test_standalone_long_digits_redacted(self):
        out, n = ipp.redact_names("Roster: 1234567 enrolled late.")
        assert "1234567" not in out
        assert n >= 1

    def test_year_not_treated_as_id(self):
        out, n = ipp.redact_names("Written in 2024 for the new unit.")
        assert "2024" in out
        # Only 4 digits — should NOT be flagged as student ID.

    # Pass 5: accommodation / IEP / 504
    def test_iep_redacted(self):
        out, n = ipp.redact_names(
            "Recall that IEP goals include reading fluency by Q3.",
        )
        assert "IEP" not in out
        assert "[ACCOMMODATION REDACTED]" in out
        assert n >= 1

    def test_504_plan_redacted(self):
        out, n = ipp.redact_names(
            "Per 504 plan, preferential seating near the board.",
        )
        assert "504 plan" not in out
        assert "[ACCOMMODATION REDACTED]" in out

    def test_speech_therapy_redacted(self):
        out, n = ipp.redact_names(
            "Pulled for speech therapy on Tuesdays during fourth period.",
        )
        assert "speech therapy" not in out
        assert "[ACCOMMODATION REDACTED]" in out

    # Pass 6: lone first names in bullet lines
    def test_lone_first_name_in_bullet_redacted(self):
        text = "Call on:\n- Marcus\n- Devon\n- Priya"
        out, n = ipp.redact_names(text)
        for name in ("Marcus", "Devon", "Priya"):
            assert name not in out, f"{name!r} leaked through"
        assert n >= 3

    def test_lone_allowlist_first_name_preserved(self):
        text = "Readings:\n- Rosa speaks out\n- Abraham addresses the nation"
        out, _ = ipp.redact_names(text)
        # Rosa and Abraham are first names from the default allowlist.
        assert "Rosa" in out
        assert "Abraham" in out

    def test_numbered_list_first_name_redacted(self):
        text = "Order:\n1. Marcus\n2. Priya"
        out, n = ipp.redact_names(text)
        assert "Marcus" not in out and "Priya" not in out
        assert n >= 2

    def test_short_lone_words_skipped(self):
        # The regex requires 3+ lowercase letters after the capital, so "A" or
        # "Hi" in a bullet line should not be redacted.
        text = "Notes:\n- Hi everyone"
        out, n = ipp.redact_names(text)
        assert out == text and n == 0

    # --- All-caps regression coverage -------------------------------------
    # The pre-fix redactor only matched title-case names. These tests lock in
    # that uppercase student names are now caught and cannot leak into the
    # voice profile (signature_phrases) as trigrams.

    def test_uppercase_full_name_redacted(self):
        out, n = ipp.redact_names("JOHN DOE answered first.")
        assert "JOHN DOE" not in out
        assert "[STUDENT]" in out
        assert n >= 1

    def test_keyword_anchored_uppercase_redacted(self):
        # This is the exact bypass reported: "STUDENT: JOHN DOE" slipped
        # through the title-case regex and leaked into trigrams.
        out, n = ipp.redact_names("STUDENT: JOHN DOE leads discussion")
        assert "JOHN" not in out
        assert "DOE" not in out
        assert "STUDENT:" not in out  # whole keyword+name replaced with [STUDENT]
        assert "[STUDENT]" in out
        assert n >= 1

    def test_keyword_anchored_titlecase_redacted(self):
        out, n = ipp.redact_names("Student: John Doe presented.")
        assert "John Doe" not in out
        assert "Student:" not in out
        assert "[STUDENT]" in out
        assert n >= 1

    def test_keyword_anchored_learner_variant_redacted(self):
        out, n = ipp.redact_names("Learner: Maria Rodriguez arrived late.")
        assert "Maria" not in out and "Rodriguez" not in out
        assert "[STUDENT]" in out
        assert n >= 1

    def test_keyword_anchored_cannot_be_bypassed_by_teacher_name(self):
        # Even when "Jane Doe" is the teacher_name (a legitimate bare-name
        # allowlist), the "Student:" / "Learner:" prefix must still cause
        # redaction — the keyword itself is the PII signal.
        out, n = ipp.redact_names(
            "Student: Jane Doe needs extra support.",
            teacher_name="Jane Doe",
        )
        assert "Jane Doe" not in out
        assert "[STUDENT]" in out
        assert n >= 1

    def test_uppercase_historical_allowlist_preserved(self):
        # All-caps allowlist comparison is case-insensitive so historical
        # figures quoted in uppercase (e.g. headings) aren't mis-redacted.
        out, n = ipp.redact_names("We discussed ROSA PARKS today.")
        assert "ROSA PARKS" in out
        assert n == 0

    def test_uppercase_structural_heading_preserved(self):
        # Template headings like "LEARNING INTENTION" and "SUCCESS CRITERIA"
        # are Cap+Cap word pairs but must not be redacted as names.
        text = "LEARNING INTENTION\nSUCCESS CRITERIA\nEXIT TICKET"
        out, n = ipp.redact_names(text)
        assert out == text
        assert n == 0

    def test_uppercase_lone_bullet_name_redacted(self):
        text = "Groups:\n- MARCUS\n- PRIYA\n- DEVON"
        out, n = ipp.redact_names(text)
        for name in ("MARCUS", "PRIYA", "DEVON"):
            assert name not in out
        assert n >= 3

    def test_uppercase_lone_bullet_structural_preserved(self):
        # Heading-style uppercase tokens in a bullet list must not be
        # treated as student first names.
        text = "Outline:\n- TODAY\n- EXIT\n- AGENDA"
        out, n = ipp.redact_names(text)
        for tok in ("TODAY", "EXIT", "AGENDA"):
            assert tok in out
        assert n == 0

    def test_uppercase_name_does_not_leak_to_signature_phrases(self):
        # Integration-style: run the full voice extractor over text that
        # previously produced "student john doe" / "john doe leads"
        # trigrams. After the fix, neither name token may appear.
        text = (
            "STUDENT: JOHN DOE leads discussion today.\n"
            "STUDENT: JOHN DOE reviews notes.\n"
            "STUDENT: JOHN DOE closes out.\n"
        ) * 2
        redacted, n = ipp.redact_names(text)
        assert n >= 1
        sig = ipp.extract_voice_signature([(redacted, 0, 0)])
        joined = " | ".join(sig["signature_phrases"])
        assert "john" not in joined.lower()
        assert "doe" not in joined.lower()


# ---------------------------------------------------------------------------
# extract_layout_signature
# ---------------------------------------------------------------------------

class TestExtractLayoutSignature:
    def test_basic_counts(self):
        # Two docs: one with 2 tables/10 rows, one with 0 tables/0 rows.
        data = [
            ("- one\n- two\nParagraph line", 2, 10),
            ("plain prose only", 0, 0),
        ]
        sig = ipp.extract_layout_signature(data)
        assert sig["avg_tables_per_doc"] == 1.0
        assert sig["avg_rows_per_table"] == 5.0
        assert 0 < sig["bullet_ratio"] <= 1

    def test_empty_data(self):
        sig = ipp.extract_layout_signature([])
        assert sig["avg_tables_per_doc"] == 0
        assert sig["avg_rows_per_table"] == 0
        assert sig["bullet_ratio"] == 0

    def test_heading_markers_counted(self):
        text = "INTRODUCTION\nregular text\nDISCUSSION\n"
        sig = ipp.extract_layout_signature([(text, 0, 0)])
        assert sig["heading_markers_count"] >= 2


# ---------------------------------------------------------------------------
# extract_voice_signature
# ---------------------------------------------------------------------------

class TestExtractVoiceSignature:
    def test_empty_returns_zeros(self):
        sig = ipp.extract_voice_signature([])
        assert sig["contraction_rate_per_1k"] == 0
        assert sig["second_person_rate_per_1k"] == 0
        assert sig["signature_phrases"] == []
        assert sig["warmth_markers_per_1k"] == 0

    def test_contractions_counted(self):
        # Heavily use contractions so the per-1k rate is clearly non-zero.
        text = "don't won't can't it's that's you're I'll we'll they'll there's"
        sig = ipp.extract_voice_signature([(text, 0, 0)])
        assert sig["contraction_rate_per_1k"] > 0

    def test_second_person_counted(self):
        text = "you will practice. your notes. students will share."
        sig = ipp.extract_voice_signature([(text, 0, 0)])
        assert sig["second_person_rate_per_1k"] > 0

    def test_numbered_bullets_counted(self):
        # Regression for audit D4: numbered bullets now count toward median.
        text = "1. Warm up activity\n2. Main task\n3. Closing reflection\n"
        sig = ipp.extract_voice_signature([(text, 0, 0)])
        assert sig["script_vs_outline_median_words"] > 0

    def test_warmth_parenthetical_counted(self):
        text = "We read aloud (groan) and then shared (bonus). " * 5
        sig = ipp.extract_voice_signature([(text, 0, 0)])
        assert sig["warmth_markers_per_1k"] > 0

    def test_signature_phrases_require_repetition(self):
        text = (
            "quick write and share " * 3 +
            "one off phrase here only once"
        )
        sig = ipp.extract_voice_signature([(text, 0, 0)])
        # The thrice-repeated phrase should show up; the once-only phrase
        # is below the >=2 threshold.
        joined = " | ".join(sig["signature_phrases"])
        assert "quick write" in joined or "write and share" in joined


# ---------------------------------------------------------------------------
# extract_activity_library
# ---------------------------------------------------------------------------

class TestExtractActivityLibrary:
    def test_recurring_threshold(self):
        docs = [
            ("We used gallery walk and stations today.", 0, 0),
            ("Another gallery walk in period 3.", 0, 0),
            ("Just a do-now, nothing fancy here.", 0, 0),
        ]
        sources = ["a.docx", "b.docx", "c.docx"]
        lib = ipp.extract_activity_library(docs, sources)
        # gallery walk appears in 2 docs → kept.
        assert "gallery walk" in lib
        assert lib["gallery walk"]["count"] == 2
        assert set(lib["gallery walk"]["sources"]) == {"a.docx", "b.docx"}
        # stations only in 1 → dropped.
        assert "stations" not in lib

    def test_case_insensitive(self):
        docs = [
            ("EXIT TICKET at end.", 0, 0),
            ("exit ticket to close.", 0, 0),
        ]
        lib = ipp.extract_activity_library(docs, ["a.docx", "b.docx"])
        assert "exit ticket" in lib


# ---------------------------------------------------------------------------
# extract_pacing
# ---------------------------------------------------------------------------

class TestExtractPacing:
    def test_iso_date_extracted(self):
        docs = [("2024-09-03 Photosynthesis unit opens today.", 0, 0)]
        pacing = ipp.extract_pacing(docs, ["p1.docx"])
        assert pacing["skipped"] == 0
        assert len(pacing["entries"]) == 1
        assert pacing["entries"][0]["date"] == "2024-09-03"

    def test_skipped_when_no_date(self):
        docs = [("No date anywhere in this plan.", 0, 0)]
        pacing = ipp.extract_pacing(docs, ["p1.docx"])
        assert pacing["skipped"] == 1
        assert pacing["entries"] == []

    def test_entries_sorted_by_date(self):
        docs = [
            ("2024-09-03 later lesson", 0, 0),
            ("2024-08-15 earlier lesson", 0, 0),
        ]
        pacing = ipp.extract_pacing(docs, ["a.docx", "b.docx"])
        dates = [e["date"] for e in pacing["entries"]]
        assert dates == sorted(dates)

    def test_week_of_format(self):
        docs = [("Week of September 9\nUnit kickoff", 0, 0)]
        pacing = ipp.extract_pacing(docs, ["w.docx"])
        assert pacing["skipped"] == 0
        assert "September" in pacing["entries"][0]["date"]


# ---------------------------------------------------------------------------
# format_markdown_output
# ---------------------------------------------------------------------------

class TestFormatMarkdownOutput:
    def _make_dims(self):
        layout = {
            "avg_tables_per_doc": 1.2, "avg_rows_per_table": 4.5,
            "bullet_ratio": 0.33, "heading_markers_count": 3,
        }
        voice = {
            "script_vs_outline_median_words": 6.0,
            "contraction_rate_per_1k": 12.34,
            "second_person_rate_per_1k": 5.0,
            "signature_phrases": ["think pair share", "gallery walk protocol"],
            "warmth_markers_per_1k": 2.1,
        }
        activities = {
            "gallery walk": {"count": 3, "sources": ["a.docx", "b.docx"]},
            "exit ticket":  {"count": 2, "sources": ["c.docx"]},
        }
        pacing = {
            "entries": [{"date": "2024-09-03", "topic": "Unit 1 kickoff"}],
            "skipped": 1,
        }
        return layout, voice, activities, pacing

    def test_frontmatter_and_sections(self):
        layout, voice, activities, pacing = self._make_dims()
        md = ipp.format_markdown_output(
            layout, voice, activities, pacing,
            source_count=4, source_files=["a.docx", "b.docx"], redactions=7,
        )
        assert md.startswith("---\n")
        assert "source_count: 4" in md
        assert "redactions: 7" in md
        assert "## Voice signature" in md
        assert "## Layout signature" in md
        assert "## Activity library" in md
        assert "## Pacing (inferred)" in md
        # Activity row rendered.
        assert "| gallery walk | 3" in md
        # Pacing entry rendered.
        assert "2024-09-03 → Unit 1 kickoff" in md
        # Skipped-date footnote rendered.
        assert "1 plan(s) had no parseable date" in md

    def test_small_sample_warning(self):
        layout, voice, activities, pacing = self._make_dims()
        md = ipp.format_markdown_output(
            layout, voice, activities, pacing,
            source_count=2, source_files=["a.docx"], redactions=0,
        )
        assert "Only 2 plans analyzed" in md

    def test_sufficient_sample_no_warning(self):
        layout, voice, activities, pacing = self._make_dims()
        md = ipp.format_markdown_output(
            layout, voice, activities, pacing,
            source_count=5, source_files=["a.docx"], redactions=0,
        )
        assert "reliable match expected" in md


# ---------------------------------------------------------------------------
# CLI — end-to-end against synthetic .docx input dir
# ---------------------------------------------------------------------------

def _seed_input_dir(input_dir: Path) -> None:
    """Populate an input directory with three synthetic past-plan .docx files."""
    input_dir.mkdir(parents=True, exist_ok=True)
    _build_docx(
        input_dir / "2024-09-03_plan.docx",
        paragraphs=[
            "2024-09-03 Photosynthesis intro",
            "Do-now: gallery walk of diagrams",
            "- Students will annotate",
            "- Think-pair-share to close",
            "John Smith joined late.",
        ],
        tables=[[["Time", "Activity"], ["5 min", "Exit ticket"]]],
    )
    _build_docx(
        input_dir / "2024-09-10_plan.docx",
        paragraphs=[
            "2024-09-10 Cell respiration",
            "Gallery walk continues today",
            "- Review ATP",
            "- Exit ticket at end",
            "Student #1234567 completed makeup work.",
        ],
    )
    _build_docx(
        input_dir / "2024-09-17_plan.docx",
        paragraphs=[
            "2024-09-17 Enzymes",
            "Warm-up: think-pair-share",
            "Exit ticket: 3 questions",
        ],
    )


class TestCli:
    def test_end_to_end_generates_profile(self, tmp_path):
        input_dir = tmp_path / "plans"
        _seed_input_dir(input_dir)
        output = tmp_path / "voice-profile.md"
        result = subprocess.run(
            [sys.executable, str(SCRIPT_PATH),
             "--input-dir", str(input_dir),
             "--output", str(output)],
            capture_output=True, text=True,
        )
        assert result.returncode == 0, result.stderr
        assert "✓" not in result.stdout
        assert "OK:" in result.stdout
        assert output.exists()
        assert output.with_suffix(".json").exists()
        md = output.read_text()
        # Redaction happened.
        assert "John Smith" not in md
        assert "1234567" not in md
        # Recurring activities identified.
        assert "gallery walk" in md or "exit ticket" in md
        # Pacing entries included.
        assert "2024-09-03" in md

    def test_json_sidecar_emitted(self, tmp_path):
        input_dir = tmp_path / "plans"
        _seed_input_dir(input_dir)
        output = tmp_path / "voice-profile.md"
        result = subprocess.run(
            [sys.executable, str(SCRIPT_PATH),
             "--input-dir", str(input_dir),
             "--output", str(output),
             "--json-sidecar"],
            capture_output=True, text=True,
        )
        assert result.returncode == 0, result.stderr
        sidecar = output.with_suffix(".json")
        assert sidecar.exists()
        data = json.loads(sidecar.read_text())
        # Dense short keys.
        assert "L" in data and "V" in data
        assert data["n"] == 3

    def test_no_json_sidecar_skips_sidecar_output(self, tmp_path):
        input_dir = tmp_path / "plans"
        _seed_input_dir(input_dir)
        output = tmp_path / "voice-profile.md"
        result = subprocess.run(
            [sys.executable, str(SCRIPT_PATH),
             "--input-dir", str(input_dir),
             "--output", str(output),
             "--no-json-sidecar"],
            capture_output=True, text=True,
        )
        assert result.returncode == 0, result.stderr
        assert output.exists()
        assert not output.with_suffix(".json").exists()

    def test_missing_input_dir(self, tmp_path):
        output = tmp_path / "out.md"
        result = subprocess.run(
            [sys.executable, str(SCRIPT_PATH),
             "--input-dir", str(tmp_path / "nope"),
             "--output", str(output)],
            capture_output=True, text=True,
        )
        assert result.returncode == 1
        assert "not found" in result.stderr.lower()

    def test_empty_input_dir(self, tmp_path):
        input_dir = tmp_path / "empty"
        input_dir.mkdir()
        output = tmp_path / "out.md"
        result = subprocess.run(
            [sys.executable, str(SCRIPT_PATH),
             "--input-dir", str(input_dir),
             "--output", str(output)],
            capture_output=True, text=True,
        )
        assert result.returncode == 1
        assert ".docx" in result.stderr.lower()

    def test_min_plans_warning(self, tmp_path):
        # Only one plan present — min-plans default of 3 should trigger a
        # warning on stderr but still succeed.
        input_dir = tmp_path / "plans"
        input_dir.mkdir()
        _build_docx(
            input_dir / "lone.docx",
            paragraphs=["2024-09-03 solo plan", "Warm-up: quick write"],
        )
        output = tmp_path / "out.md"
        result = subprocess.run(
            [sys.executable, str(SCRIPT_PATH),
             "--input-dir", str(input_dir),
             "--output", str(output)],
            capture_output=True, text=True,
        )
        assert result.returncode == 0
        assert "recommended" in result.stderr.lower()

    def test_no_redact_flag(self, tmp_path):
        input_dir = tmp_path / "plans"
        input_dir.mkdir()
        _build_docx(
            input_dir / "plan.docx",
            paragraphs=[
                "2024-09-03 opener",
                "John Smith answered first.",
                "Gallery walk then exit ticket.",
                "Also gallery walk tomorrow.",
            ],
        )
        # Add a second doc so extract_activity_library has >=2 occurrences.
        _build_docx(
            input_dir / "plan2.docx",
            paragraphs=[
                "2024-09-04 follow-up",
                "Gallery walk to review.",
                "Exit ticket at the end.",
            ],
        )
        output = tmp_path / "out.md"
        result = subprocess.run(
            [sys.executable, str(SCRIPT_PATH),
             "--input-dir", str(input_dir),
             "--output", str(output),
             "--no-redact-names"],
            capture_output=True, text=True,
        )
        assert result.returncode == 0, result.stderr
        md = output.read_text()
        # The analysis output never dumps raw plan text verbatim, so we can't
        # assert "John Smith" in md; but the redaction counter should be 0.
        assert "redactions: 0" in md

    def test_teacher_name_allowlisted_via_cli(self, tmp_path):
        input_dir = tmp_path / "plans"
        input_dir.mkdir()
        _build_docx(
            input_dir / "p1.docx",
            paragraphs=[
                "2024-09-03 lesson",
                "Taught by Jane Doe this week.",
                "Gallery walk activity.",
            ],
        )
        _build_docx(
            input_dir / "p2.docx",
            paragraphs=[
                "2024-09-10 lesson",
                "Jane Doe introduces the unit.",
                "Gallery walk continues.",
            ],
        )
        output = tmp_path / "out.md"
        result = subprocess.run(
            [sys.executable, str(SCRIPT_PATH),
             "--input-dir", str(input_dir),
             "--output", str(output),
             "--teacher-name", "Jane Doe"],
            capture_output=True, text=True,
        )
        assert result.returncode == 0, result.stderr
        # Teacher name should not have been redacted (though it won't appear
        # in the summary markdown by default — we check the redaction count).
        assert "redactions: 0" in output.read_text()

    def test_uppercase_student_name_does_not_leak_to_profile(self, tmp_path):
        # End-to-end regression for the reported release blocker: three
        # .docx inputs whose text contains "STUDENT: JOHN DOE" must not
        # leak "john doe" into the generated voice-profile.md (either as
        # raw text or as trigrams in signature_phrases).
        input_dir = tmp_path / "plans"
        input_dir.mkdir()
        for i in range(1, 4):
            _build_docx(
                input_dir / f"plan_{i}.docx",
                paragraphs=[
                    f"2024-09-0{i} lesson",
                    "STUDENT: JOHN DOE leads discussion today.",
                    "- Warm-up",
                    "Gallery walk activity.",
                ],
            )
        output = tmp_path / "voice-profile.md"
        result = subprocess.run(
            [sys.executable, str(SCRIPT_PATH),
             "--input-dir", str(input_dir),
             "--output", str(output)],
            capture_output=True, text=True,
        )
        assert result.returncode == 0, result.stderr
        md = output.read_text()
        md_lower = md.lower()
        assert "john doe" not in md_lower
        assert "john" not in md_lower, "JOHN token leaked into profile"
        # Redaction counter must be non-zero; the pre-fix run reported 0.
        assert "redactions: 0" not in md

    def test_metadata_authors_are_redacted_before_profile_analysis(self, tmp_path):
        input_dir = tmp_path / "plans"
        input_dir.mkdir()
        docx_path = input_dir / "plan.docx"

        doc = Document()
        para = doc.add_paragraph("2024-09-03 Photosynthesis intro")
        run = para.add_run(" Gallery walk to open.")
        doc.core_properties.author = "Taylor Johnson"
        doc.core_properties.last_modified_by = "Taylor Johnson"
        doc.add_comment(run, text="Looks good", author="Maria Rodriguez")
        doc.save(str(docx_path))
        _inject_tracked_change_author(docx_path, "Alex Kim")

        output = tmp_path / "voice-profile.md"
        result = subprocess.run(
            [sys.executable, str(SCRIPT_PATH),
             "--input-dir", str(input_dir),
             "--output", str(output)],
            capture_output=True, text=True,
        )

        assert result.returncode == 0, result.stderr
        md = output.read_text()
        assert "Taylor Johnson" not in md
        assert "Maria Rodriguez" not in md
        assert "Alex Kim" not in md
        assert "redactions: 3" in md
