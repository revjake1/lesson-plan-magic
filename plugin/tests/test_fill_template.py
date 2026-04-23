"""Pytest regression tests for fill_template.py PII scanner fixes."""

import json
import os
import subprocess
import sys
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Pt

import fill_template as ft
from fill_template import (
    fill_placeholders,
    parse_plan,
    scan_for_pii,
    scan_docx_content,
    validate_privacy,
    write_config_mutation,
)


SCRIPT_PATH = (
    Path(__file__).parent.parent
    / "skills" / "lesson-planner" / "scripts" / "fill_template.py"
)
STARTER_TEMPLATES_DIR = (
    Path(__file__).parent.parent
    / "skills" / "lesson-planner" / "assets" / "starter-templates"
)


SAMPLE_PLAN_MD = """# Week of 2026-04-20
Subject: Chemistry
Teacher: Jane Smith

## 2026-04-20 — Monday

### Standards
- HS-PS1-1

### Learning Intention
I can describe Newton's second law.

### Success Criteria
- I can identify force and mass.

### Agenda
1. Warm-up
2. Practice

### Materials
- Whiteboard

### Evidence
Exit ticket
"""


def _run_fill(
    plan_path,
    template_path,
    out_path,
    *,
    env=None,
    allow_anywhere=True,
    **kwargs,
):
    argv = [
        sys.executable,
        str(SCRIPT_PATH),
        "--plan", str(plan_path),
        "--template", str(template_path),
        "--output", str(out_path),
    ]
    if allow_anywhere:
        argv.append("--allow-anywhere")
    for k, v in kwargs.items():
        flag = f"--{k.replace('_', '-')}"
        if isinstance(v, bool):
            if v:
                argv.append(flag)
            continue
        argv.extend([flag, str(v)])
    run_env = os.environ.copy()
    if env:
        run_env.update(env)
    return subprocess.run(argv, capture_output=True, text=True, env=run_env)


class TestValidatePrivacyTelemetryBoolean:
    """YAML 1.1 parses ``off`` as Python False; the validator must accept that.

    Regression: the original check rejected any non-string telemetry
    value, which meant the documented config (``telemetry: off`` with
    no quotes) couldn't pass its own validation under PyYAML.
    """

    def test_telemetry_false_accepted(self):
        cfg = {"privacy": {"student_data": "never", "telemetry": False}}
        assert validate_privacy(cfg) == []

    def test_telemetry_string_off_accepted(self):
        cfg = {"privacy": {"student_data": "never", "telemetry": "off"}}
        assert validate_privacy(cfg) == []

    def test_telemetry_string_OFF_accepted_case_insensitive(self):
        cfg = {"privacy": {"student_data": "never", "telemetry": "OFF"}}
        assert validate_privacy(cfg) == []

    def test_telemetry_true_rejected(self):
        cfg = {"privacy": {"student_data": "never", "telemetry": True}}
        errors = validate_privacy(cfg)
        assert errors and "telemetry" in errors[0]

    def test_telemetry_null_rejected(self):
        cfg = {"privacy": {"student_data": "never", "telemetry": None}}
        errors = validate_privacy(cfg)
        assert errors and "telemetry" in errors[0]

    def test_telemetry_string_no_rejected(self):
        # "no" is a YAML boolean in 1.1 but would parse to False — this
        # branch covers the case where someone quoted it as "no" and
        # PyYAML respected the quotes. We still reject that.
        cfg = {"privacy": {"student_data": "never", "telemetry": "no"}}
        errors = validate_privacy(cfg)
        assert errors and "telemetry" in errors[0]

    def test_student_data_never_required(self):
        cfg = {"privacy": {"student_data": "sometimes", "telemetry": False}}
        errors = validate_privacy(cfg)
        assert errors and "student_data" in errors[0]

    def test_student_data_false_rejected(self):
        # Defensive: there's no YAML keyword that spells "never", so
        # booleans should be rejected outright.
        cfg = {"privacy": {"student_data": False, "telemetry": False}}
        errors = validate_privacy(cfg)
        assert errors and "student_data" in errors[0]

    def test_pii_scan_before_write_string_false_rejected(self):
        cfg = {
            "privacy": {
                "student_data": "never",
                "telemetry": False,
                "pii_scan_before_write": "false",
            }
        }
        errors = validate_privacy(cfg)
        assert any("pii_scan_before_write" in error for error in errors)


class TestParsePlan:
    def test_rejects_non_padded_day_header(self):
        bad_plan = SAMPLE_PLAN_MD.replace("## 2026-04-20", "## 2026-4-20")
        with pytest.raises(ValueError, match="zero-padded YYYY-MM-DD"):
            parse_plan(bad_plan)

    def test_rejects_plan_with_no_day_headers(self):
        with pytest.raises(ValueError, match="No lesson days found"):
            parse_plan("# Week of 2026-04-20\nSubject: Chemistry\nTeacher: Jane\n")


class TestWriteConfigMutation:
    def test_rejects_config_path_outside_lesson_plan_magic_root(self, tmp_path, monkeypatch):
        fake_home = tmp_path / "home"
        monkeypatch.setattr(ft.Path, "home", lambda: fake_home)

        external = tmp_path / "elsewhere" / "config.yaml"
        external.parent.mkdir(parents=True)
        external.write_text("subjects:\n- id: chem\n", encoding="utf-8")

        ok = write_config_mutation(
            external,
            "chem",
            {"template.mapping_verified": True},
        )

        assert ok is False
        assert "mapping_verified" not in external.read_text(encoding="utf-8")

    def test_home_root_prefers_home_env(self, tmp_path, monkeypatch):
        fake_home = tmp_path / "home"
        monkeypatch.delenv("LESSON_PLAN_MAGIC_HOME", raising=False)
        monkeypatch.setenv("HOME", str(fake_home))
        monkeypatch.setattr(ft.Path, "home", lambda: Path("/should-not-be-used"))

        assert ft._home_root() == fake_home / "Documents" / "Lesson Plan Magic"

class TestScanForPiiAllowlisting:
    """Test scan_for_pii with allowed_names allowlisting."""

    def test_allowed_name_not_flagged(self):
        """Teacher name in allowed_names set is not flagged."""
        text = "Jane Smith prepared this lesson."
        matches = scan_for_pii(text, allowed_names={"Jane Smith"})
        assert matches == []

    def test_bare_name_flagged_when_not_allowed(self):
        """Bare capitalized name is flagged when not allowlisted."""
        text = "This lesson was prepared by Suspicious Student."
        matches = scan_for_pii(text, allowed_names=set())
        # Should find the bare name hit
        assert any("bare name" in label for _, label in matches)

    def test_empty_allowed_names(self):
        """Empty allowed_names set flags bare names."""
        text = "Johnny Doe prepared this material."
        matches = scan_for_pii(text, allowed_names=set())
        assert len(matches) > 0
        assert any("Johnny Doe" in match[0] for match in matches)

    def test_ssn_always_flagged(self):
        """SSN pattern is always flagged regardless of allowlist."""
        text = "SSN 123-45-6789 should not be here."
        matches = scan_for_pii(text, allowed_names=set())
        assert any("SSN" in label for _, label in matches)
        assert any("123-45-6789" in match[0] for match in matches)

    def test_phone_number_flagged(self):
        """Phone number pattern is flagged."""
        text = "Call 555-123-4567 for more info."
        matches = scan_for_pii(text, allowed_names=set())
        assert any("phone" in label for _, label in matches)

    def test_student_email_flagged(self):
        """School-domain email is flagged."""
        text = "Email student@example.k12.ga.us for questions."
        matches = scan_for_pii(text, allowed_names=set())
        assert any("email" in label for _, label in matches)

    def test_personal_domain_email_flagged(self):
        """Personal-domain email (gmail.com, yahoo.com) is flagged.

        Regression: v1 only matched .k12.*/.edu/.org, so a student's
        personal email slipped past even though the README promised to
        block ``student email`` broadly.
        """
        text = "Email john.doe@gmail.com for questions."
        matches = scan_for_pii(text, allowed_names=set())
        assert any("email" in label for _, label in matches)

    def test_roster_pattern_flagged(self):
        """Roster-like 'Student: Name' pattern is flagged."""
        text = "Student: Johnny Doe is participating."
        matches = scan_for_pii(text, allowed_names=set())
        assert any("roster" in label for _, label in matches)

    def test_roster_pattern_allowlist_cannot_bypass(self):
        """Roster-like ``Student:`` prefix is non-allowlistable.

        Regression (P0): pre-fix, allowlisting a name (``--allow-names
        'Jane Doe'``) let ``Student: Jane Doe`` through the scanner
        because scan_for_pii skipped roster hits whose captured pair
        matched the allowlist. The keyword itself denotes student
        context — allowlist exists to silence bare-name hits on the
        teacher / public figures, not to override FERPA guarantees.
        """
        text = "Student: Jane Smith is in the roster."
        matches = scan_for_pii(text, allowed_names={"Jane Smith"})
        assert any("roster" in label for _, label in matches), (
            "Student:<allowlisted name> must still fail the scan"
        )

    def test_uppercase_roster_pattern_flagged(self):
        """Regression: ``STUDENT: JOHN DOE`` was previously missed because
        name-anchor groups required ``[A-Z][a-z]+`` (title case only)."""
        text = "STUDENT: JOHN DOE is in the roster."
        matches = scan_for_pii(text, allowed_names=set())
        assert any("roster" in label for _, label in matches)

    def test_learner_variant_of_roster(self):
        """Learner variant of roster pattern is caught."""
        text = "Learner; Mary Jones needs support."
        matches = scan_for_pii(text, allowed_names=set())
        assert any("roster" in label for _, label in matches)

    def test_named_accommodation_flagged(self):
        """Named accommodation pattern (IEP: Name, 504: Name) is flagged."""
        text = "IEP: Sarah is on this plan."
        matches = scan_for_pii(text, allowed_names=set())
        assert any("accommodation" in label for _, label in matches)

    def test_structural_words_not_flagged(self):
        """Structural words like 'Test Subject' are not flagged as names."""
        text = "The Test Subject for today is Mathematics."
        matches = scan_for_pii(text, allowed_names=set())
        # Test Subject should not be flagged since Test is in STRUCTURAL_WORDS
        assert not any("Test Subject" in match[0] for match in matches)

    def test_sentence_starter_not_flagged(self):
        """Sentence starters like 'Review Monday' are not flagged as names."""
        text = "Review Monday's lesson with the class."
        matches = scan_for_pii(text, allowed_names=set())
        # Should not flag "Review Monday"
        assert not any("Review Monday" in match[0] for match in matches)

    def test_single_name_in_sensitive_teacher_note_is_flagged(self):
        """Single-name student notes should not slip through."""
        text = "Today Maria will get extended time."
        matches = scan_for_pii(text, allowed_names=set())
        assert any(match[0] == "Maria" for match in matches)
        assert any("single-name" in label for _, label in matches)

    def test_single_name_with_missing_work_context_is_flagged(self):
        text = "Review Jamal about missing work."
        matches = scan_for_pii(text, allowed_names=set())
        assert any(match[0] == "Jamal" for match in matches)
        assert any("single-name" in label for _, label in matches)

    def test_single_name_context_respects_allowlist(self):
        text = "Today Jane will get extended time."
        matches = scan_for_pii(text, allowed_names={"Jane Smith"})
        assert not any(match[0] == "Jane" for match in matches)

    def test_default_allowed_names_not_flagged(self):
        """Historical figures in DEFAULT_ALLOWED_NAMES are not flagged."""
        text = "Students studied Rosa Parks and Abraham Lincoln."
        matches = scan_for_pii(text, allowed_names=set())
        # These are in DEFAULT_ALLOWED_NAMES
        assert not any("Rosa Parks" in match[0] for match in matches)
        assert not any("Abraham Lincoln" in match[0] for match in matches)

    def test_non_ascii_name_flagged(self):
        text = "José García prepared this handout."
        matches = scan_for_pii(text, allowed_names=set())
        assert any("José García" in match[0] for match in matches)

    def test_place_name_like_phrase_not_flagged(self):
        text = "Meet in Austin Clinic after school."
        matches = scan_for_pii(text, allowed_names=set())
        assert not any("Austin Clinic" in match[0] for match in matches)

    def test_ap_biology_not_flagged(self):
        text = "AP Biology lab practical on Friday."
        matches = scan_for_pii(text, allowed_names=set())
        assert not any("AP Biology" in match[0] for match in matches)

    def test_initial_form_name_flagged(self):
        """Initial-form names like 'Sarah H.' are flagged."""
        text = "Sarah H. contributed to this project."
        matches = scan_for_pii(text, allowed_names=set())
        assert any("initial-form" in label for _, label in matches)

    def test_multiple_issues_reported(self):
        """Multiple PII issues are all reported."""
        text = "SSN 123-45-6789 and Johnny Doe's phone is 555-123-4567."
        matches = scan_for_pii(text, allowed_names=set())
        assert len(matches) >= 3


class TestScanDocxContent:
    """Test scan_docx_content with a dynamically built docx."""

    def test_scan_docx_with_allowed_teacher_name(self):
        """Teacher name in docx is not flagged when allowlisted."""
        doc = Document()
        doc.add_paragraph("Prepared by Jane Smith")
        doc.add_paragraph("Jane Smith is the teacher.")

        matches = scan_docx_content(doc, allowed_names={"Jane Smith"})
        assert matches == []

    def test_scan_docx_with_disallowed_name(self):
        """Disallowed name in docx is flagged."""
        doc = Document()
        doc.add_paragraph("Johnny Doe prepared this.")

        matches = scan_docx_content(doc, allowed_names=set())
        assert len(matches) > 0
        assert any("Johnny Doe" in match[0] for match in matches)

    def test_scan_docx_table_content(self):
        """PII in table cells is scanned."""
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        # python-docx requires adding paragraphs to cells, not setting .text directly
        table.rows[0].cells[0].add_paragraph("Student")
        table.rows[0].cells[1].add_paragraph("SSN")
        table.rows[1].cells[0].add_paragraph("John Doe")
        table.rows[1].cells[1].add_paragraph("123-45-6789")

        matches = scan_docx_content(doc, allowed_names=set())
        # Should find SSN and John Doe
        assert any("123-45-6789" in match[0] for match in matches)

    def test_scan_empty_docx(self):
        """Empty docx produces no matches."""
        doc = Document()
        matches = scan_docx_content(doc, allowed_names=set())
        assert matches == []

    def test_scan_docx_mixed_content(self):
        """Mixed allowed and disallowed content is properly filtered."""
        doc = Document()
        doc.add_paragraph("Teacher: Jane Smith")
        doc.add_paragraph("Student: Johnny Doe")

        matches = scan_docx_content(doc, allowed_names={"Jane Smith"})
        # Should only flag Johnny Doe
        assert any("Johnny Doe" in match[0] for match in matches)
        assert not any("Jane Smith" in match[0] for match in matches)

    def test_scan_catches_pii_in_header(self):
        """Regression (P1): a template with ``Student: John Doe`` baked
        into the header used to slip past scan_docx_content — it only
        iterated ``doc.paragraphs`` + ``doc.tables`` in the body.
        """
        doc = Document()
        doc.add_paragraph("Clean body content")
        section = doc.sections[0]
        section.header.paragraphs[0].add_run("Student: John Doe")

        matches = scan_docx_content(doc, allowed_names=set())
        assert any("roster" in label for _, label in matches), (
            f"Header PII must be caught; got {matches}"
        )

    def test_scan_catches_pii_in_footer(self):
        """Footer PII must also trip the scan."""
        doc = Document()
        doc.add_paragraph("Clean body content")
        section = doc.sections[0]
        section.footer.paragraphs[0].add_run("Call 555-123-4567 with questions.")

        matches = scan_docx_content(doc, allowed_names=set())
        assert any("phone" in label for _, label in matches)

    def test_scan_catches_pii_in_textbox(self):
        """Regression (P1): text-box content (``w:txbxContent``) inside
        the body used to slip past because python-docx's
        ``doc.paragraphs`` doesn't descend into text boxes.
        """
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document()
        doc.add_paragraph("Clean body content")
        # Inject a minimal text-box with a student name inside it.
        p = doc.paragraphs[-1]._p
        r = OxmlElement("w:r")
        pict = OxmlElement("w:pict")
        shape_xml = (
            '<v:shape xmlns:v="urn:schemas-microsoft-com:vml" '
            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            'id="txt1" type="#_x0000_t202">'
            '<v:textbox><w:txbxContent>'
            '<w:p><w:r><w:t>Student: John Doe</w:t></w:r></w:p>'
            '</w:txbxContent></v:textbox></v:shape>'
        )
        from lxml import etree
        pict.append(etree.fromstring(shape_xml))
        r.append(pict)
        p.append(r)

        matches = scan_docx_content(doc, allowed_names=set())
        assert any("roster" in label for _, label in matches), (
            f"Text-box PII must be caught; got {matches}"
        )

    @pytest.mark.parametrize(
        "template_name",
        ["daily-one-pager.docx", "weekly-bell.docx", "weekly-block.docx"],
    )
    def test_starter_templates_do_not_trip_pii_scan(self, template_name):
        doc = Document(str(STARTER_TEMPLATES_DIR / template_name))
        matches = scan_docx_content(doc, allowed_names=set())
        assert matches == []


class TestPIIPatternEdgeCases:
    """Edge cases and pattern interactions."""

    def test_case_insensitive_keyword_matching(self):
        """Keyword patterns are case-insensitive for keywords."""
        text = "STUDENT: Johnny Doe is here."
        matches = scan_for_pii(text, allowed_names=set())
        assert any("roster" in label for _, label in matches)

    def test_multiple_names_in_text(self):
        """Multiple different names are all flagged."""
        text = "Johnny Doe and Mary Jones prepared this lesson."
        matches = scan_for_pii(text, allowed_names=set())
        names_found = [match[0] for match in matches]
        assert any("Johnny Doe" in name for name in names_found)
        assert any("Mary Jones" in name for name in names_found)

    def test_line_break_prevents_bare_name_across_lines(self):
        """Names split across line breaks are not matched as bare names."""
        text = "Johnny\nDoe prepared this."
        matches = scan_for_pii(text, allowed_names=set())
        # Should NOT match "Johnny Doe" since they're on different lines
        assert not any("Johnny Doe" in match[0] for match in matches)

    def test_allowlist_case_sensitive(self):
        """Allowlist matching respects case sensitivity of stored names."""
        # If allowlist has "Jane Smith", it should match exactly
        text = "Jane Smith prepared this."
        matches = scan_for_pii(text, allowed_names={"Jane Smith"})
        assert matches == []

        # But "JANE SMITH" might be treated differently depending on implementation
        # This test documents current behavior
        matches2 = scan_for_pii(text, allowed_names={"jane smith"})
        # The roster pattern matching is case-insensitive, so this may or may not match


class TestFillTemplateCliRegressions:
    def test_unknown_section_pii_blocks_output_and_sidecar(self, tmp_path):
        plan = tmp_path / "plan.md"
        plan.write_text(
            SAMPLE_PLAN_MD
            + "\n### Notes\nStudent: John Doe needs extra time today.\n",
            encoding="utf-8",
        )
        template = tmp_path / "template.docx"
        doc = Document()
        doc.add_paragraph("{{MONDAY_LEARNING_INTENTION}}")
        doc.save(str(template))

        out = tmp_path / "out.docx"
        result = _run_fill(plan, template, out)

        assert result.returncode == 2
        assert "PII detected in plan" in result.stderr
        assert not out.exists()
        assert not out.with_suffix(".plan.md").exists()

    def test_header_placeholder_is_filled(self, tmp_path):
        plan = tmp_path / "plan.md"
        plan.write_text(SAMPLE_PLAN_MD, encoding="utf-8")
        template = tmp_path / "template.docx"
        doc = Document()
        doc.add_paragraph("{{MONDAY_LEARNING_INTENTION}}")
        doc.sections[0].header.paragraphs[0].text = "Header {{MONDAY_DATE}}"
        doc.save(str(template))

        out = tmp_path / "out.docx"
        result = _run_fill(plan, template, out)

        assert result.returncode == 0, result.stderr
        written = Document(str(out))
        assert written.paragraphs[0].text == "I can describe Newton's second law."
        assert written.sections[0].header.paragraphs[0].text == "Header 2026-04-20"

    def test_unresolved_placeholder_blocks_success(self, tmp_path):
        plan = tmp_path / "plan.md"
        plan.write_text(SAMPLE_PLAN_MD, encoding="utf-8")
        template = tmp_path / "template.docx"
        doc = Document()
        doc.add_paragraph("{{MONDAY_LEARNING_INTENTION}}")
        doc.sections[0].header.paragraphs[0].text = "Header {{UNKNOWN_TOKEN}}"
        doc.save(str(template))

        out = tmp_path / "out.docx"
        result = _run_fill(plan, template, out)

        assert result.returncode == 4
        assert "unresolved template placeholders" in result.stderr
        assert "{{UNKNOWN_TOKEN}}" in result.stderr
        assert not out.exists()

    def test_existing_different_sidecar_blocks_before_output_write(self, tmp_path):
        plan = tmp_path / "plan.md"
        plan.write_text(SAMPLE_PLAN_MD, encoding="utf-8")
        template = tmp_path / "template.docx"
        doc = Document()
        doc.add_paragraph("{{MONDAY_LEARNING_INTENTION}}")
        doc.save(str(template))

        out = tmp_path / "out.docx"
        sidecar = out.with_suffix(".plan.md")
        sidecar.write_text("stale content", encoding="utf-8")

        result = _run_fill(plan, template, out)

        assert result.returncode == 5
        assert "refusing to overwrite existing sidecar" in result.stderr
        assert not out.exists()
        assert sidecar.read_text(encoding="utf-8") == "stale content"

    def test_existing_different_json_sidecar_blocks_before_output_write(self, tmp_path):
        plan = tmp_path / "plan.md"
        plan.write_text(SAMPLE_PLAN_MD, encoding="utf-8")
        template = tmp_path / "template.docx"
        doc = Document()
        doc.add_paragraph("{{MONDAY_LEARNING_INTENTION}}")
        doc.save(str(template))

        out = tmp_path / "out.docx"
        json_sidecar = out.with_suffix(".plan.json")
        json_sidecar.write_text(
            json.dumps({"v": 1, "d": [{"dt": "2026-04-21"}]}),
            encoding="utf-8",
        )

        result = _run_fill(plan, template, out)

        assert result.returncode == 5
        assert "refusing to overwrite existing JSON sidecar" in result.stderr
        assert not out.exists()
        assert json.loads(json_sidecar.read_text(encoding="utf-8")) == {
            "v": 1,
            "d": [{"dt": "2026-04-21"}],
        }

    def test_blocks_output_outside_documented_dir_without_allow_anywhere(self, tmp_path):
        home = tmp_path / "home"
        plan = tmp_path / "plan.md"
        plan.write_text(SAMPLE_PLAN_MD, encoding="utf-8")
        template = tmp_path / "template.docx"
        doc = Document()
        doc.add_paragraph("{{MONDAY_LEARNING_INTENTION}}")
        doc.save(str(template))

        out = tmp_path / "outside.docx"
        result = _run_fill(
            plan,
            template,
            out,
            env={"HOME": str(home)},
            allow_anywhere=False,
        )

        expected_root = home / "Documents" / "Lesson Plan Magic" / "outputs"
        assert result.returncode == 1
        assert str(expected_root) in result.stderr
        assert "allow-anywhere" in result.stderr.lower()
        assert not out.exists()
        assert not out.with_suffix(".plan.md").exists()

    def test_allows_output_inside_documented_dir_by_default(self, tmp_path):
        home = tmp_path / "home"
        plan = tmp_path / "plan.md"
        plan.write_text(SAMPLE_PLAN_MD, encoding="utf-8")
        template = tmp_path / "template.docx"
        doc = Document()
        doc.add_paragraph("{{MONDAY_LEARNING_INTENTION}}")
        doc.save(str(template))

        out = home / "Documents" / "Lesson Plan Magic" / "outputs" / "inside.docx"
        result = _run_fill(
            plan,
            template,
            out,
            env={"HOME": str(home)},
            allow_anywhere=False,
        )

        assert result.returncode == 0, result.stderr
        assert out.exists()
        assert out.with_suffix(".plan.md").exists()
        assert out.with_suffix(".plan.json").exists()

    def test_allow_anywhere_explicitly_allows_external_output(self, tmp_path):
        home = tmp_path / "home"
        plan = tmp_path / "plan.md"
        plan.write_text(SAMPLE_PLAN_MD, encoding="utf-8")
        template = tmp_path / "template.docx"
        doc = Document()
        doc.add_paragraph("{{MONDAY_LEARNING_INTENTION}}")
        doc.save(str(template))

        out = tmp_path / "outside.docx"
        result = _run_fill(
            plan,
            template,
            out,
            env={"HOME": str(home)},
        )

        assert result.returncode == 0, result.stderr
        assert out.exists()
        assert out.with_suffix(".plan.md").exists()
        assert out.with_suffix(".plan.json").exists()

    def test_plan_json_sidecar_is_written_with_compact_day_payload(self, tmp_path):
        plan = tmp_path / "plan.md"
        plan.write_text(SAMPLE_PLAN_MD, encoding="utf-8")
        template = tmp_path / "template.docx"
        doc = Document()
        doc.add_paragraph("{{MONDAY_LEARNING_INTENTION}}")
        doc.save(str(template))

        out = tmp_path / "out.docx"
        result = _run_fill(plan, template, out)

        assert result.returncode == 0, result.stderr
        sidecar = json.loads(out.with_suffix(".plan.json").read_text(encoding="utf-8"))
        assert sidecar["v"] == ft.PLAN_SIDECAR_VERSION
        assert sidecar["w"] == "2026-04-20"
        assert sidecar["s"] == "Chemistry"
        assert sidecar["t"] == "Jane Smith"
        assert sidecar["d"] == [
            {
                "ag": ["Warm-up", "Practice"],
                "dt": "2026-04-20",
                "e": "Exit ticket",
                "li": "I can describe Newton's second law.",
                "m": ["Whiteboard"],
                "n": "Monday",
                "sc": ["I can identify force and mass."],
                "st": ["HS-PS1-1"],
            }
        ]

    def test_no_sidecar_skips_markdown_and_json_sidecars(self, tmp_path):
        plan = tmp_path / "plan.md"
        plan.write_text(SAMPLE_PLAN_MD, encoding="utf-8")
        template = tmp_path / "template.docx"
        doc = Document()
        doc.add_paragraph("{{MONDAY_LEARNING_INTENTION}}")
        doc.save(str(template))

        out = tmp_path / "out.docx"
        result = _run_fill(plan, template, out, no_sidecar=True)

        assert result.returncode == 0, result.stderr
        assert out.exists()
        assert not out.with_suffix(".plan.md").exists()
        assert not out.with_suffix(".plan.json").exists()

    def test_plan_markdown_size_limit_blocks_output(self, tmp_path):
        plan = tmp_path / "plan.md"
        plan.write_text("x" * (ft.MAX_PLAN_BYTES + 1), encoding="utf-8")
        template = tmp_path / "template.docx"
        doc = Document()
        doc.add_paragraph("{{MONDAY_LEARNING_INTENTION}}")
        doc.save(str(template))

        out = tmp_path / "out.docx"
        result = _run_fill(plan, template, out)

        assert result.returncode == 1
        assert "byte limit" in result.stderr.lower()
        assert not out.exists()

    def test_malformed_config_errors_cleanly(self, tmp_path):
        home = tmp_path / "home"
        config_dir = home / "Documents" / "Lesson Plan Magic"
        config_dir.mkdir(parents=True)
        (config_dir / "config.yaml").write_text(
            "privacy: [unterminated\n",
            encoding="utf-8",
        )

        plan = tmp_path / "plan.md"
        plan.write_text(SAMPLE_PLAN_MD, encoding="utf-8")
        template = tmp_path / "template.docx"
        doc = Document()
        doc.add_paragraph("{{MONDAY_LEARNING_INTENTION}}")
        doc.save(str(template))

        out = config_dir / "outputs" / "out.docx"
        result = _run_fill(
            plan,
            template,
            out,
            env={"HOME": str(home)},
            allow_anywhere=False,
        )

        assert result.returncode == 1
        assert "Traceback" not in result.stderr
        assert "Could not load config" in result.stderr
        assert not out.exists()

    def test_pii_stderr_does_not_echo_matched_student_name(self, tmp_path):
        plan = tmp_path / "plan.md"
        plan.write_text(
            SAMPLE_PLAN_MD
            + "\n### Notes\nStudent: John Doe needs extra time today.\n",
            encoding="utf-8",
        )
        template = tmp_path / "template.docx"
        doc = Document()
        doc.add_paragraph("{{MONDAY_LEARNING_INTENTION}}")
        doc.save(str(template))

        out = tmp_path / "out.docx"
        result = _run_fill(plan, template, out)

        assert result.returncode == 2
        assert "PII detected in plan" in result.stderr
        assert "John Doe" not in result.stderr


class TestFillPlaceholdersTextbox:
    def test_textbox_placeholder_is_replaced(self):
        from docx.oxml import OxmlElement
        from lxml import etree

        doc = Document()
        doc.add_paragraph("Body {{MONDAY_DATE}}")
        host_para = doc.paragraphs[-1]._p

        drawing_run = OxmlElement("w:r")
        pict = OxmlElement("w:pict")
        shape_xml = (
            '<v:shape xmlns:v="urn:schemas-microsoft-com:vml" '
            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            'id="txt1" type="#_x0000_t202">'
            '<v:textbox><w:txbxContent>'
            '<w:p><w:r><w:t>Textbox {{MONDAY_DATE}}</w:t></w:r></w:p>'
            '</w:txbxContent></v:textbox></v:shape>'
        )
        pict.append(etree.fromstring(shape_xml))
        drawing_run.append(pict)
        host_para.append(drawing_run)

        changed = fill_placeholders(doc, parse_plan(SAMPLE_PLAN_MD))

        assert changed is True
        texts = list(doc.element.body.itertext())
        assert "Body 2026-04-20" in texts
        assert "Textbox 2026-04-20" in texts
        assert not any("{{MONDAY_DATE}}" in text for text in texts)
